from __future__ import annotations

from io import BytesIO
import logging
from pathlib import Path
from time import perf_counter
import re
from uuid import UUID, uuid4

from docx import Document
import numpy as np
from PIL import Image, ImageColor, ImageDraw, ImageFont, ImageOps
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.contracts import GeneratedImageAsset, GenerationSceneGraph, RendererInput, RendererResponse, SceneGraphElement
from app.core.config import get_settings
from app.integrations.object_storage import LocalObjectStorage
from app.utils.image_assets import open_image_asset
from app.utils.palette_roles import derive_palette_roles

logger = logging.getLogger(__name__)


class RendererService:
    SCENE_TEXT_SYMBOL_PATTERN = re.compile(r"(?:âœ”ï¸\x8f|âœ”ï¸|âœ…|â€¢|✔️|✓|✅|➡️|➜)")

    def __init__(self, session: AsyncSession) -> None:
        self.session = session
        self.storage = LocalObjectStorage()
        self.settings = get_settings()
        self.payload: RendererInput | None = None
        self._active_font_candidates: list[Path] = []
        self._active_font_bindings: list[dict[str, str]] = []
        self._used_font_paths: set[str] = set()
        self._used_font_families: set[str] = set()
        self._requested_font_families: set[str] = set()

    @staticmethod
    def _is_bold_weight(weight: str | None) -> bool:
        normalized = str(weight or "").strip().casefold().replace("_", "").replace("-", "").replace(" ", "")
        return normalized in {"bold", "semibold", "demibold", "extrabold", "black", "heavy", "700", "800", "900"}

    def _font_candidates(self, weight: str | None = None) -> list[Path]:
        repo_root = Path(__file__).resolve().parents[2]
        configured = Path(self.settings.renderer_font_path)
        weighted_candidates: list[Path] = []
        if self._is_bold_weight(weight):
            weighted_candidates = [
                repo_root / "frontend" / "public" / "fonts" / "DM_Sans" / "static" / "DMSans-Bold.ttf",
                repo_root / "frontend" / "public" / "fonts" / "DM_Sans" / "static" / "DMSans-SemiBold.ttf",
                repo_root / "frontend" / "public" / "fonts" / "Manrope" / "static" / "Manrope-Bold.ttf",
                repo_root / "frontend" / "public" / "fonts" / "Manrope" / "static" / "Manrope-SemiBold.ttf",
            ]
        candidates = [
            *self._active_font_candidates,
            *weighted_candidates,
            configured,
            repo_root / "frontend" / "public" / "fonts" / "DM_Sans" / "static" / "DMSans-Regular.ttf",
            repo_root / "frontend" / "public" / "fonts" / "Manrope" / "static" / "Manrope-Regular.ttf",
        ]
        deduped: list[Path] = []
        seen: set[str] = set()
        for candidate in candidates:
            key = str(candidate)
            if key in seen:
                continue
            seen.add(key)
            deduped.append(candidate)
        return deduped

    def _font_path_for_family(self, family_hint: str | None) -> Path | None:
        normalized_hint = str(family_hint or "").strip().casefold()
        if not normalized_hint:
            return None
        self._requested_font_families.add(str(family_hint).strip())
        for binding in self._active_font_bindings:
            family_name = str(binding.get("family_name") or "").strip()
            storage_path = str(binding.get("storage_path") or "").strip()
            if not family_name or not storage_path:
                continue
            if family_name.casefold() == normalized_hint:
                return Path(self.storage.absolute_path(storage_path))
        return None

    def _font(self, size: int, family_hint: str | None = None, weight: str | None = None) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
        preferred = self._font_path_for_family(family_hint)
        if preferred and preferred.exists():
            try:
                self._used_font_paths.add(str(preferred))
                self._used_font_families.add(str(family_hint or "").strip())
                return ImageFont.truetype(str(preferred), size=size)
            except OSError:
                pass
        for font_path in self._font_candidates(weight=weight):
            if font_path.exists():
                try:
                    self._used_font_paths.add(str(font_path))
                    return ImageFont.truetype(str(font_path), size=size)
                except OSError:
                    continue
        return ImageFont.load_default()

    @staticmethod
    def _zone_box(zone) -> tuple[int, int, int, int]:
        return zone.x, zone.y, zone.x + zone.width, zone.y + zone.height

    def _parse_color(self, value: str | None, default: tuple[int, int, int]) -> tuple[int, int, int]:
        if not value:
            return default
        try:
            return ImageColor.getrgb(value)
        except ValueError:
            return default

    @staticmethod
    def _palette_role_map(palette: object) -> dict[str, str]:
        if isinstance(palette, dict):
            return {
                str(key).strip().lower(): str(value).strip()
                for key, value in palette.items()
                if str(value).strip()
            }
        role_map: dict[str, str] = {}
        if not isinstance(palette, list):
            return role_map
        for entry in palette:
            if not isinstance(entry, dict):
                continue
            role = str(entry.get("role") or "").strip().lower()
            value = str(
                entry.get("hex_code")
                or entry.get("hex")
                or entry.get("value")
                or entry.get("color_code")
                or ""
            ).strip()
            if role and value:
                role_map[role] = value
        return role_map

    @staticmethod
    def _style_color_payload_to_hex(payload: dict | None) -> str | None:
        if not isinstance(payload, dict):
            return None
        value = str(payload.get("hex_code") or payload.get("hex") or "").strip()
        return value or None

    @staticmethod
    def _composition_plan(payload: RendererInput) -> dict:
        if not payload.blueprint:
            return {}
        return payload.blueprint.composition_plan or {}

    @staticmethod
    def _should_render_scene_graph_direct(payload: RendererInput) -> bool:
        if not payload.scene_graph or not payload.scene_graph.elements:
            return False
        if payload.base_canvas_asset_path:
            return True
        format_name = str(payload.studio_panel.get("format") or payload.blueprint.layout_type or "static").strip().lower()
        file_type = str(payload.studio_panel.get("file_type") or payload.blueprint.export_format or "png").strip().lower()
        source_mode = str(payload.blueprint.source_mode or "").strip().lower()
        if format_name == "infographic" and source_mode in {"exact_template", "adapted_template"}:
            return True
        if format_name in {"carousel", "pdf"}:
            return False
        if file_type in {"pdf", "doc"}:
            return False
        return True

    @staticmethod
    def _prefer_blueprint_zone_rendering(payload: RendererInput) -> bool:
        if not payload.blueprint or not payload.blueprint.zones:
            return False
        source_mode = str(payload.blueprint.source_mode or "").strip().lower()
        if source_mode not in {"exact_template", "adapted_template"}:
            return False
        roles = {
            str(zone.role or "").strip().lower()
            for zone in payload.blueprint.zones
            if getattr(zone, "role", None)
        }
        return bool(roles & {"logo", "headline", "body", "cta", "image"})

    def _template_style_reference(self, payload: RendererInput) -> dict:
        template_intelligence = payload.brand_visual_rules.get("template_intelligence", []) if payload.brand_visual_rules else []
        if not isinstance(template_intelligence, list):
            return {}
        source_template_id = str(payload.blueprint.source_template_id or "").strip()
        for template in template_intelligence:
            if not isinstance(template, dict):
                continue
            if source_template_id and str(template.get("template_id", "")).strip() == source_template_id:
                return template.get("analysis", template) if isinstance(template.get("analysis"), dict) else template
        for template in template_intelligence:
            if not isinstance(template, dict):
                continue
            return template.get("analysis", template) if isinstance(template.get("analysis"), dict) else template
        return {}

    def _footer_text(self, payload: RendererInput, page_text: dict | None = None) -> str:
        page_footer = str((page_text or {}).get("footer") or "").strip()
        if page_footer:
            return page_footer
        metadata = payload.text.metadata if isinstance(payload.text.metadata, dict) else {}
        metadata_footer = str(metadata.get("footer") or "").strip()
        if metadata_footer:
            return metadata_footer
        template_ref = self._template_style_reference(payload)
        template_footer = str(template_ref.get("footer") or "").strip() if isinstance(template_ref, dict) else ""
        if template_footer:
            return template_footer
        return ""

    def _zone_style_hint(self, payload: RendererInput, role: str) -> dict | None:
        template_ref = self._template_style_reference(payload)
        if not template_ref:
            return None
        direct_map = {
            "headline": template_ref.get("heading_style"),
            "header": template_ref.get("header_style"),
            "footer": template_ref.get("footer_style"),
        }
        if isinstance(direct_map.get(role), dict):
            return direct_map[role]
        text_style_map = template_ref.get("text_style_map", [])
        if not isinstance(text_style_map, list):
            return None
        if role == "body":
            excluded = {
                str(item.get("text", "")).strip().casefold()
                for item in [template_ref.get("heading_style"), template_ref.get("header_style"), template_ref.get("footer_style")]
                if isinstance(item, dict) and item.get("text")
            }
            for item in text_style_map:
                if not isinstance(item, dict):
                    continue
                text = str(item.get("text", "")).strip()
                if not text or text.casefold() in excluded:
                    continue
                if len(text.split()) >= 4:
                    return item
            return next((item for item in text_style_map if isinstance(item, dict)), None)
        if role == "cta":
            cta_tokens = ("apply", "learn more", "download", "get started", "register", "book", "join")
            for item in text_style_map:
                if not isinstance(item, dict):
                    continue
                lowered = str(item.get("text", "")).strip().lower()
                if any(token in lowered for token in cta_tokens):
                    return item
        return None

    def _resolved_fill_from_style_hint(
        self,
        style_hint: dict | None,
        default: tuple[int, int, int],
    ) -> tuple[int, int, int]:
        if not isinstance(style_hint, dict):
            return default
        return self._parse_color(self._style_color_payload_to_hex(style_hint.get("font_color")), default)

    def _resolve_gradient_spec(
        self,
        payload: RendererInput,
        background: tuple[int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
    ) -> dict | None:
        composition = self._composition_plan(payload)
        background_plan = composition.get("background_plan", {}) if isinstance(composition, dict) else {}
        if background_plan.get("policy") == "template_background":
            return None
        if background_plan.get("policy") == "brand_solid_background":
            return None

        palette = self._resolve_palette_roles(payload)
        balance = self._reference_color_balance(payload, palette)
        if balance:
            bg_share = float(balance.get("background_share") or 0.0)
            accent_share = float(balance.get("accent_share") or 0.0)
            primary_share = float(balance.get("primary_share") or 0.0)
            background_hex = self._normalize_hex(balance.get("background_hex") if isinstance(balance, dict) else None)
            primary_hex = self._normalize_hex(balance.get("primary_hex") if isinstance(balance, dict) else None)
            accent_hex = self._normalize_hex(balance.get("accent_hex") if isinstance(balance, dict) else None)
            if background_hex and primary_hex and accent_hex:
                if bg_share >= 0.58 and accent_share <= 0.14:
                    return {
                        "type": "linear",
                        "direction": "vertical",
                        "start_color": background_hex,
                        "mid_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(self._parse_color(background_hex, background), self._parse_color(primary_hex, primary), 0.05 + (primary_share * 0.18))),
                        "end_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(self._parse_color(background_hex, background), self._parse_color(accent_hex, accent), 0.03 + (accent_share * 0.2))),
                        "confidence": 0.82,
                        "source": "reference_color_usage",
                    }
                if accent_share >= 0.2:
                    return {
                        "type": "linear",
                        "direction": "vertical",
                        "start_color": background_hex,
                        "mid_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(self._parse_color(background_hex, background), self._parse_color(accent_hex, accent), 0.08 + (accent_share * 0.28))),
                        "end_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(self._parse_color(background_hex, background), self._parse_color(primary_hex, primary), 0.08 + (primary_share * 0.22))),
                        "confidence": 0.79,
                        "source": "reference_color_usage",
                    }

        gradient_preferences = payload.brand_visual_rules.get("gradient_preferences") if payload.brand_visual_rules else None
        if isinstance(gradient_preferences, list):
            for item in gradient_preferences:
                if isinstance(item, dict) and item.get("start_color") and item.get("end_color"):
                    return item
        if isinstance(gradient_preferences, dict) and gradient_preferences.get("start_color") and gradient_preferences.get("end_color"):
            return gradient_preferences

        template_ref = self._template_style_reference(payload)
        gradients = template_ref.get("gradients", []) if isinstance(template_ref, dict) else []
        for item in gradients if isinstance(gradients, list) else []:
            if isinstance(item, dict) and item.get("start_color") and item.get("end_color"):
                return item

        if payload.blueprint.source_mode == "exact_template":
            return None
        return {
            "type": "linear",
            "direction": "vertical",
            "start_color": "#{:02X}{:02X}{:02X}".format(*background),
            "end_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(background, accent, 0.22)),
            "mid_color": "#{:02X}{:02X}{:02X}".format(*self._blend_color(background, primary, 0.12)),
            "confidence": 0.58,
            "source": "brand_palette_fallback",
        }

    def _build_background_canvas(
        self,
        width: int,
        height: int,
        background: tuple[int, int, int],
        gradient_spec: dict | None,
        base_canvas_asset_path: str | None = None,
    ) -> Image.Image:
        if base_canvas_asset_path:
            source = Path(self.storage.absolute_path(base_canvas_asset_path))
            if source.exists():
                try:
                    with open_image_asset(source) as raw:
                        return self._fit_image_to_zone(raw.convert("RGB"), width, height)
                except OSError:
                    logger.warning("renderer.base_canvas.unreadable storage_path=%s", base_canvas_asset_path)
        if not gradient_spec:
            return Image.new("RGB", (width, height), color=background)

        gradient_type = str(gradient_spec.get("type", "linear")).lower()
        direction = str(gradient_spec.get("direction", "vertical")).lower()
        start = np.array(self._parse_color(str(gradient_spec.get("start_color") or ""), background), dtype=np.float32)
        end = np.array(self._parse_color(str(gradient_spec.get("end_color") or ""), background), dtype=np.float32)
        mid_color = np.array(
            self._parse_color(str(gradient_spec.get("mid_color") or ""), tuple(((start + end) / 2).astype(int))),
            dtype=np.float32,
        )
        if gradient_type == "radial":
            x = np.linspace(-1.0, 1.0, width, dtype=np.float32)
            y = np.linspace(-1.0, 1.0, height, dtype=np.float32)
            xx, yy = np.meshgrid(x, y)
            distance = np.sqrt((xx ** 2) + (yy ** 2))
            ratio = np.clip(distance / np.max(distance), 0.0, 1.0)[:, :, None]
            color_array = (mid_color * (1 - ratio)) + (end * ratio)
        else:
            if direction == "horizontal":
                ratio = np.linspace(0.0, 1.0, width, dtype=np.float32)[None, :, None]
                color_array = (start * (1 - ratio)) + (end * ratio)
                color_array = np.repeat(color_array, height, axis=0)
            else:
                ratio = np.linspace(0.0, 1.0, height, dtype=np.float32)[:, None, None]
                color_array = (start * (1 - ratio)) + (end * ratio)
                color_array = np.repeat(color_array, width, axis=1)
        return Image.fromarray(np.clip(color_array, 0, 255).astype(np.uint8), mode="RGB")

    def _metadata_palette(self, payload: RendererInput) -> dict[str, str]:
        colors = payload.text.metadata.get("brand_colors") or payload.text.metadata.get("colors") or {}
        if not isinstance(colors, dict):
            return {}
        return {str(key): str(value) for key, value in colors.items()}

    def _reference_color_usage_entries(self, payload: RendererInput) -> list[dict[str, object]]:
        template_zone_map = (payload.template_metadata or {}).get("zone_map", {}) if payload.template_metadata else {}
        template_ref = self._template_style_reference(payload)
        sources = [
            template_zone_map.get("color_usage") if isinstance(template_zone_map, dict) else None,
            template_ref.get("color_usage") if isinstance(template_ref, dict) else None,
            payload.brand_visual_rules.get("color_usage") if payload.brand_visual_rules else None,
        ]
        best_source: list[dict[str, object]] = []
        for candidate in sources:
            if not isinstance(candidate, list):
                continue
            normalized_source: list[dict[str, object]] = []
            for entry in candidate:
                if not isinstance(entry, dict):
                    continue
                hex_code = self._normalize_hex(
                    entry.get("hex_code")
                    or entry.get("hex")
                    or entry.get("value")
                    or entry.get("color_code")
                )
                if not hex_code:
                    continue
                try:
                    count = float(entry.get("count") or entry.get("weight") or entry.get("ratio") or 1.0)
                except (TypeError, ValueError):
                    count = 1.0
                normalized_source.append(
                    {
                        "hex_code": hex_code,
                        "role": str(entry.get("role") or "").strip().lower(),
                        "count": max(count, 0.0),
                    }
                )
            if len(normalized_source) > len(best_source):
                best_source = normalized_source
        return best_source

    def _reference_color_balance(
        self,
        payload: RendererInput,
        palette: dict[str, str],
    ) -> dict[str, object]:
        entries = self._reference_color_usage_entries(payload)
        if not entries:
            return {}
        merged: dict[str, dict[str, object]] = {}
        for entry in entries:
            hex_code = str(entry.get("hex_code") or "")
            role = str(entry.get("role") or "").strip().lower()
            count = float(entry.get("count") or 0.0)
            current = merged.setdefault(hex_code, {"hex_code": hex_code, "roles": set(), "count": 0.0})
            current["count"] = float(current.get("count") or 0.0) + count
            if role:
                current["roles"].add(role)
        normalized_entries = sorted(
            [
                {
                    "hex_code": key,
                    "roles": sorted(value.get("roles") or []),
                    "count": float(value.get("count") or 0.0),
                }
                for key, value in merged.items()
            ],
            key=lambda item: float(item.get("count") or 0.0),
            reverse=True,
        )
        total = sum(float(item.get("count") or 0.0) for item in normalized_entries)
        if total <= 0:
            return {}

        def _pick_hex(role_tokens: set[str]) -> str | None:
            for item in normalized_entries:
                roles = {str(role).strip().lower() for role in (item.get("roles") or [])}
                if roles & role_tokens:
                    return str(item.get("hex_code") or "")
            return None

        background_hex = _pick_hex({"background", "surface", "canvas"})
        primary_hex = _pick_hex({"primary", "brand", "headline", "title"})
        accent_hex = _pick_hex({"accent", "secondary", "highlight", "cta"})

        if not background_hex:
            light_entries = [
                item
                for item in normalized_entries
                if (rgb := self._hex_to_rgb(str(item.get("hex_code") or ""))) and self._relative_luminance(rgb) >= 0.72
            ]
            background_hex = str((light_entries[0] if light_entries else normalized_entries[0]).get("hex_code") or "")
        if not primary_hex:
            primary_hex = next(
                (
                    str(item.get("hex_code") or "")
                    for item in normalized_entries
                    if str(item.get("hex_code") or "") != background_hex
                ),
                palette.get("primary") or palette.get("brand") or "",
            )
        if not accent_hex:
            accent_hex = next(
                (
                    str(item.get("hex_code") or "")
                    for item in normalized_entries
                    if str(item.get("hex_code") or "") not in {background_hex, primary_hex}
                ),
                palette.get("accent") or palette.get("secondary") or primary_hex,
            )

        def _share(hex_code: str | None) -> float:
            if not hex_code:
                return 0.0
            return sum(
                float(item.get("count") or 0.0)
                for item in normalized_entries
                if str(item.get("hex_code") or "") == hex_code
            ) / total

        return {
            "entries": normalized_entries,
            "background_hex": background_hex or palette.get("background") or palette.get("surface"),
            "primary_hex": primary_hex or palette.get("primary") or palette.get("brand"),
            "accent_hex": accent_hex or palette.get("accent") or palette.get("secondary") or primary_hex,
            "background_share": _share(background_hex),
            "primary_share": _share(primary_hex),
            "accent_share": _share(accent_hex),
        }

    def _reference_palette_tokens(
        self,
        payload: RendererInput,
        *,
        palette: dict[str, str],
        background: tuple[int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
    ) -> dict[str, tuple[int, int, int]]:
        balance = self._reference_color_balance(payload, palette)
        background_share = float(balance.get("background_share") or 0.0)
        primary_share = float(balance.get("primary_share") or 0.0)
        accent_share = float(balance.get("accent_share") or 0.0)
        white = (255, 255, 255)

        background_rgb = self._hex_to_rgb(str(balance.get("background_hex") or "")) or background
        primary_rgb = self._hex_to_rgb(str(balance.get("primary_hex") or "")) or primary
        accent_rgb = self._hex_to_rgb(str(balance.get("accent_hex") or "")) or accent

        bg_is_light = self._relative_luminance(background_rgb) >= 0.72
        surface_alpha = 0.18 if background_share >= 0.5 else 0.32
        if not bg_is_light:
            surface_alpha = 0.1 if background_share >= 0.5 else 0.18
        surface = (
            self._blend_color(background_rgb, white, surface_alpha)
            if bg_is_light
            else self._blend_color(background_rgb, white, surface_alpha)
        )
        accent_alpha = max(0.06, min(0.22, 0.04 + (accent_share * 0.7)))
        primary_alpha = max(0.05, min(0.2, 0.05 + (primary_share * 0.45)))
        stroke_alpha = max(0.1, min(0.24, 0.1 + (primary_share * 0.35)))
        proof_surface_alpha = 0.38 if background_share >= 0.5 else 0.54
        frame_alpha = max(0.04, min(0.16, 0.03 + (primary_share * 0.28)))

        return {
            "background": background_rgb,
            "primary": primary_rgb,
            "accent": accent_rgb,
            "surface": surface,
            "accent_soft": self._blend_color(background_rgb, accent_rgb, accent_alpha),
            "primary_soft": self._blend_color(background_rgb, primary_rgb, primary_alpha),
            "stroke": self._blend_color(background_rgb, primary_rgb, stroke_alpha),
            "proof_surface": self._blend_color(background_rgb, surface, proof_surface_alpha),
            "image_frame_fill": self._blend_color(background_rgb, primary_rgb, frame_alpha),
        }

    @staticmethod
    def _hex_to_rgb(value: str) -> tuple[int, int, int] | None:
        text = str(value or "").strip()
        if not re.fullmatch(r"#?[0-9a-fA-F]{6}", text):
            return None
        normalized = text[1:] if text.startswith("#") else text
        return tuple(int(normalized[index:index + 2], 16) for index in range(0, 6, 2))

    def _inferred_palette_roles(self, payload: RendererInput) -> dict[str, str]:
        return derive_palette_roles(payload.brand_visual_rules or {})

    def _resolve_palette_roles(self, payload: RendererInput) -> dict[str, str]:
        rules_palette = self._inferred_palette_roles(payload)
        allowed_values = {
            self._normalize_hex(value)
            for value in rules_palette.values()
            if self._normalize_hex(value)
        }
        metadata_palette: dict[str, str] = {}
        for key, value in self._metadata_palette(payload).items():
            normalized_key = str(key).strip().lower()
            normalized_value = self._normalize_hex(value)
            if not normalized_value:
                continue
            if allowed_values:
                if normalized_key in {"background", "surface", "canvas"}:
                    rgb = self._hex_to_rgb(normalized_value)
                    if normalized_value not in allowed_values and not (rgb and self._is_soft_neutral_color(rgb)):
                        continue
                elif normalized_value not in allowed_values:
                    continue
            metadata_palette[normalized_key] = normalized_value
        return {**metadata_palette, **rules_palette}

    @staticmethod
    def _normalize_hex(value: str | None) -> str | None:
        text = str(value or "").strip().upper()
        if not text:
            return None
        if not text.startswith("#") and re.fullmatch(r"[0-9A-F]{6}", text):
            text = f"#{text}"
        return text if re.fullmatch(r"#[0-9A-F]{6}", text) else None

    @staticmethod
    def _is_soft_neutral_color(rgb: tuple[int, int, int]) -> bool:
        return max(rgb) - min(rgb) <= 24 or sum(rgb) >= 660

    def _resolve_primary_color(self, payload: RendererInput, palette: dict[str, str]) -> tuple[int, int, int]:
        balance = self._reference_color_balance(payload, palette)
        reference_primary = self._normalize_hex(balance.get("primary_hex") if isinstance(balance, dict) else None)
        if reference_primary:
            return self._parse_color(reference_primary, (20, 20, 20))
        return self._parse_color(
            palette.get("primary") or palette.get("brand") or palette.get("headline"),
            (20, 20, 20),
        )

    def _resolve_accent_color(
        self,
        payload: RendererInput,
        palette: dict[str, str],
        primary: tuple[int, int, int],
    ) -> tuple[int, int, int]:
        balance = self._reference_color_balance(payload, palette)
        reference_accent = self._normalize_hex(balance.get("accent_hex") if isinstance(balance, dict) else None)
        if reference_accent:
            return self._parse_color(reference_accent, primary)
        explicit = (
            palette.get("secondary")
            or palette.get("accent")
            or palette.get("highlight")
            or palette.get("cta")
        )
        if explicit:
            return self._parse_color(explicit, primary)
        return self._blend_color(primary, (255, 255, 255), 0.18)

    def _resolve_background_color(
        self,
        payload: RendererInput,
        palette: dict[str, str],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
    ) -> tuple[int, int, int]:
        # Fix 1: Check scene_graph.styles.background_fill first (highest priority)
        if payload.scene_graph and payload.scene_graph.styles:
            scene_background = payload.scene_graph.styles.get("background_fill")
            if scene_background:
                normalized = self._normalize_hex(scene_background)
                if normalized:
                    return self._parse_color(normalized, (247, 244, 236))

        balance = self._reference_color_balance(payload, palette)
        reference_background = self._normalize_hex(balance.get("background_hex") if isinstance(balance, dict) else None)
        if reference_background:
            return self._parse_color(reference_background, (247, 244, 236))
        explicit_background = payload.brand_visual_rules.get("background_color") if payload.brand_visual_rules else None
        normalized_explicit = self._normalize_hex(explicit_background)
        palette_values = {
            self._normalize_hex(value)
            for value in palette.values()
            if self._normalize_hex(value)
        }
        if normalized_explicit:
            rgb = self._hex_to_rgb(normalized_explicit)
            if normalized_explicit in palette_values or (rgb and self._is_soft_neutral_color(rgb)):
                return self._parse_color(normalized_explicit, (247, 244, 236))

        explicit = palette.get("background") or palette.get("surface") or palette.get("canvas")
        if explicit:
            return self._parse_color(explicit, (247, 244, 236))

        palette_candidates = [
            rgb
            for rgb in (self._hex_to_rgb(value) for value in palette.values())
            if rgb and (
                self._is_soft_neutral_color(rgb)
                or sum(rgb) >= 610
                or (rgb[0] >= 220 and rgb[1] >= 205 and rgb[2] <= 205)
            )
        ]
        if palette_candidates:
            return max(
                palette_candidates,
                key=lambda rgb: (sum(rgb), -abs(rgb[0] - rgb[1]), -abs(rgb[1] - rgb[2])),
            )

        # Derive a soft branded surface instead of falling back to a blue-weighted canvas.
        base = self._blend_color((255, 255, 255), primary, 0.05)
        if accent != primary:
            base = self._blend_color(base, accent, 0.03)
        return base

    def _brand_has_real_logo(self, payload: RendererInput) -> bool:
        identity = payload.brand_visual_rules.get("identity", {}) if payload.brand_visual_rules else {}
        if payload.logo_asset_path:
            return True
        if not isinstance(identity, dict):
            return False
        return bool(identity.get("logo_asset_id") or identity.get("logo_asset_ids") or identity.get("logo_assets"))

    @staticmethod
    def _template_is_style_reference_only(payload: RendererInput) -> bool:
        scene_graph_hints = payload.scene_graph.validation_hints if payload.scene_graph else {}
        if isinstance(scene_graph_hints, dict) and str(scene_graph_hints.get("template_surface_policy") or "").strip() == "style_reference_only":
            return True
        creative_decision = payload.creative_decision if isinstance(payload.creative_decision, dict) else {}
        asset_strategy = creative_decision.get("asset_strategy", {}) if isinstance(creative_decision, dict) else {}
        if isinstance(asset_strategy, dict) and str(asset_strategy.get("template_surface_policy") or "").strip() == "style_reference_only":
            return True
        return False

    def _render_logo_box(
        self,
        *,
        canvas: Image.Image,
        draw: ImageDraw.ImageDraw,
        payload: RendererInput,
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        fill: tuple[int, int, int],
        logo_path: str | None = None,
    ) -> bool:
        resolved_logo_path = logo_path or payload.logo_asset_path
        if resolved_logo_path and not self._storage_exists(resolved_logo_path):
            recovered_logo_path = self._discover_logo_storage_path(payload)
            if recovered_logo_path:
                logger.warning(
                    "renderer.logo.recovered brand_space_id=%s content_version_id=%s storage_path=%s",
                    payload.brand_space_id,
                    payload.content_version_id,
                    recovered_logo_path,
                )
                resolved_logo_path = recovered_logo_path
        zone = type("Zone", (), self._zone_manifest("logo", "logo", box))()
        if resolved_logo_path:
            return self._paste_logo(canvas, resolved_logo_path, zone)
        recovered_logo_path = self._discover_logo_storage_path(payload)
        if recovered_logo_path:
            return self._paste_logo(canvas, recovered_logo_path, zone)
        if self._brand_has_real_logo(payload):
            logger.warning(
                "renderer.logo.real_asset_unavailable brand_space_id=%s content_version_id=%s",
                payload.brand_space_id,
                payload.content_version_id,
            )
            return False
        self._draw_brand_lockup(draw, self._brand_name(payload), box, primary, accent, fill)
        return True

    def _storage_exists(self, storage_path: str | None) -> bool:
        if not storage_path:
            return False
        exists_fn = getattr(self.storage, "exists", None)
        if callable(exists_fn):
            try:
                return bool(exists_fn(storage_path))
            except Exception:
                return False
        try:
            return Path(self.storage.absolute_path(storage_path)).exists()
        except Exception:
            return False

    @staticmethod
    def _edge_background_should_strip(image: Image.Image, threshold: int = 245) -> bool:
        rgba = image.convert("RGBA")
        width, height = rgba.size
        if width <= 0 or height <= 0:
            return False
        edge_pixels: list[tuple[int, int, int, int]] = []
        pixels = rgba.load()
        for x in range(width):
            edge_pixels.append(pixels[x, 0])
            edge_pixels.append(pixels[x, height - 1])
        for y in range(1, max(height - 1, 1)):
            edge_pixels.append(pixels[0, y])
            edge_pixels.append(pixels[width - 1, y])
        opaque_edges = [pixel for pixel in edge_pixels if pixel[3] > 0]
        if not opaque_edges:
            return False
        light_edges = [
            pixel
            for pixel in opaque_edges
            if pixel[0] >= threshold and pixel[1] >= threshold and pixel[2] >= threshold
        ]
        return (len(light_edges) / len(opaque_edges)) >= 0.75

    @classmethod
    def _strip_logo_background_if_safe(cls, image: Image.Image) -> Image.Image:
        rgba = image.convert("RGBA")
        if not cls._edge_background_should_strip(rgba):
            return rgba
        width, height = rgba.size
        pixels = rgba.load()
        keep = [[True for _ in range(width)] for _ in range(height)]
        queue: list[tuple[int, int]] = []
        seen: set[tuple[int, int]] = set()

        def is_background(px: tuple[int, int, int, int]) -> bool:
            red, green, blue, alpha = px
            return alpha > 0 and red >= 245 and green >= 245 and blue >= 245

        for x in range(width):
            queue.append((x, 0))
            queue.append((x, height - 1))
        for y in range(height):
            queue.append((0, y))
            queue.append((width - 1, y))

        while queue:
            x, y = queue.pop()
            if (x, y) in seen or x < 0 or y < 0 or x >= width or y >= height:
                continue
            seen.add((x, y))
            if not is_background(pixels[x, y]):
                continue
            keep[y][x] = False
            queue.extend(((x - 1, y), (x + 1, y), (x, y - 1), (x, y + 1)))

        cleaned = rgba.copy()
        cleaned_pixels = cleaned.load()
        for y in range(height):
            for x in range(width):
                if not keep[y][x]:
                    red, green, blue, _alpha = cleaned_pixels[x, y]
                    cleaned_pixels[x, y] = (red, green, blue, 0)
        return cleaned

    @staticmethod
    def _trim_transparent_logo_margins(image: Image.Image) -> Image.Image:
        rgba = image.convert("RGBA")
        alpha = rgba.getchannel("A")
        bbox = alpha.getbbox()
        if not bbox:
            return rgba
        left, top, right, bottom = bbox
        if left == 0 and top == 0 and right == rgba.width and bottom == rgba.height:
            return rgba
        return rgba.crop(bbox)

    def _discover_logo_storage_path(self, payload: RendererInput) -> str | None:
        base_path = getattr(self.storage, "base_path", None)
        if not base_path:
            return None
        base_path = Path(base_path).resolve()
        brand_root = Path(base_path) / str(payload.tenant_id) / str(payload.brand_space_id)
        if not brand_root.exists():
            return None
        search_roots = [brand_root / "logo", brand_root / "brand", brand_root / "uploads"]
        image_suffixes = {".png", ".jpg", ".jpeg", ".webp", ".svg"}
        fallback: list[Path] = []
        for root in search_roots:
            if not root.exists():
                continue
            for candidate in sorted(root.rglob("*")):
                if not candidate.is_file() or candidate.suffix.lower() not in image_suffixes:
                    continue
                lowered_name = candidate.name.casefold()
                if any(token in lowered_name for token in ("logo", "wordmark", "brandmark", "lockup", "emblem")):
                    try:
                        return str(candidate.relative_to(base_path)).replace("\\", "/")
                    except ValueError:
                        continue
                fallback.append(candidate)
        for candidate in fallback:
            try:
                return str(candidate.relative_to(base_path)).replace("\\", "/")
            except ValueError:
                continue
        return None

    def _brand_name(self, payload: RendererInput) -> str:
        identity = payload.brand_visual_rules.get("identity", {}) if payload.brand_visual_rules else {}
        for candidate in (
            payload.brand_visual_rules.get("brand_name") if payload.brand_visual_rules else None,
            identity.get("brand_name") if isinstance(identity, dict) else None,
            payload.text.metadata.get("brand"),
        ):
            if candidate and str(candidate).strip():
                return str(candidate).strip()
        headline_words = payload.text.headline.split()
        return headline_words[0].strip().upper() if headline_words else "BRAND"

    @staticmethod
    def _badge_label(value: str) -> str:
        text = value.strip().lstrip("#").replace("_", " ").replace("-", " ")
        if not text:
            return ""
        label = " ".join(part.capitalize() for part in text.split())
        return label[:22].rstrip()

    def _supporting_badges(self, payload: RendererInput, preset: str) -> list[str]:
        raw_tags = payload.text.hashtags[:4]
        labels = [self._badge_label(tag) for tag in raw_tags]
        labels = [label for label in labels if label]
        max_badges = 3 if preset == "instagram" else 2
        return labels[:max_badges]

    def _metadata_text(self, payload: RendererInput, key: str, max_chars: int = 120) -> str:
        value = payload.text.metadata.get(key, "")
        text = " ".join(str(value).strip().split()) if value is not None else ""
        return text[:max_chars].rstrip(" ,.;:") if text else ""

    def _metadata_list(self, payload: RendererInput, key: str, limit: int) -> list[str]:
        value = payload.text.metadata.get(key, [])
        if isinstance(value, str):
            parts = re.split(r"(?:\r?\n|[;|])+|\s*•\s*|\s*[-*]\s+", value)
            if len(parts) == 1:
                parts = re.split(r",\s+|(?<=[.!?])\s+", value)
        elif isinstance(value, list):
            parts = [str(item).strip() for item in value]
        else:
            parts = []
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in parts:
            text = re.sub(r"^[#*\-\d\.\)\s]+", "", item).strip()
            if not text:
                continue
            text = " ".join(text.split())
            key_value = text.casefold()
            if key_value in seen:
                continue
            seen.add(key_value)
            cleaned.append(text[:88].rstrip(" ,.;:"))
            if len(cleaned) >= limit:
                break
        return cleaned

    def _section_label(self, payload: RendererInput, preset: str) -> str:
        label = self._metadata_text(payload, "section_label", 26)
        if label:
            return label
        badges = self._supporting_badges(payload, preset)
        if badges:
            return badges[0]
        return preset.replace("_", " ").title()

    def _supporting_line(self, payload: RendererInput, preset: str, body_text: str) -> str:
        support = self._metadata_text(payload, "supporting_line", 116) or self._metadata_text(payload, "subheadline", 116)
        if support:
            return support
        if not body_text:
            return ""
        sentences = self._split_sentences(body_text)
        if not sentences:
            return ""
        # Format-aware character limits - data-rich formats get more space
        format_name = str(payload.studio_panel.get("format") or "").strip().lower()
        if preset == "instagram":
            max_chars = 120
        elif format_name in ["infographic", "static"]:
            max_chars = 180  # More space for data-rich content
        else:
            max_chars = 144
        return self._truncate_copy_at_word_boundary(sentences[0], max_chars)

    def _proof_points(self, payload: RendererInput, preset: str, body_text: str) -> list[str]:
        points = self._metadata_list(payload, "proof_points", 4)
        if points:
            return points[: (3 if preset in {"linkedin", "infographic"} else 2)]
        # Format-aware character limits - data-rich formats get more space
        format_name = str(payload.studio_panel.get("format") or "").strip().lower()
        max_chars = 140 if format_name in ["infographic", "static"] else 108
        sentences = self._split_sentences(body_text)
        cleaned = [
            self._truncate_copy_at_word_boundary(sentence, max_chars, ellipsis=False)
            for sentence in sentences[1:4]
            if sentence.strip()
        ]
        if not cleaned and sentences:
            cleaned = [self._truncate_copy_at_word_boundary(sentences[0], max_chars, ellipsis=False)]
        return cleaned[: (3 if preset in {"linkedin", "infographic"} else 2)]

    def _stat_highlights(self, payload: RendererInput, preset: str) -> list[str]:
        stats = self._metadata_list(payload, "stat_highlights", 3)
        if stats:
            return stats[: (3 if preset != "instagram" else 2)]
        badges = self._supporting_badges(payload, preset)
        return badges[: (3 if preset != "instagram" else 2)]

    @staticmethod
    def _blend_color(
        base: tuple[int, int, int],
        overlay: tuple[int, int, int],
        alpha: float,
    ) -> tuple[int, int, int]:
        clamped = max(0.0, min(alpha, 1.0))
        return tuple(
            int(round((base[index] * (1 - clamped)) + (overlay[index] * clamped)))
            for index in range(3)
        )

    @staticmethod
    def _relative_luminance(color: tuple[int, int, int]) -> float:
        def _channel(value: int) -> float:
            normalized = max(0.0, min(value / 255.0, 1.0))
            if normalized <= 0.03928:
                return normalized / 12.92
            return ((normalized + 0.055) / 1.055) ** 2.4

        red, green, blue = (_channel(component) for component in color)
        return (0.2126 * red) + (0.7152 * green) + (0.0722 * blue)

    @classmethod
    def _contrast_ratio(
        cls,
        foreground: tuple[int, int, int],
        background: tuple[int, int, int],
    ) -> float:
        lighter = max(cls._relative_luminance(foreground), cls._relative_luminance(background))
        darker = min(cls._relative_luminance(foreground), cls._relative_luminance(background))
        return (lighter + 0.05) / (darker + 0.05)

    @classmethod
    def _auto_text_color(cls, background: tuple[int, int, int]) -> tuple[int, int, int]:
        dark = (20, 20, 20)
        light = (255, 255, 255)
        return dark if cls._contrast_ratio(dark, background) >= cls._contrast_ratio(light, background) else light

    @classmethod
    def _ensure_text_contrast(
        cls,
        foreground: tuple[int, int, int],
        background: tuple[int, int, int],
        *,
        minimum_ratio: float = 4.2,
        fallback: tuple[int, int, int] | None = None,
    ) -> tuple[int, int, int]:
        if cls._contrast_ratio(foreground, background) >= minimum_ratio:
            return foreground
        candidate = fallback or cls._auto_text_color(background)
        if cls._contrast_ratio(candidate, background) >= minimum_ratio:
            return candidate
        return cls._auto_text_color(background)

    def _wrap_text(
        self,
        draw: ImageDraw.ImageDraw,
        text: str,
        font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
        width: int,
    ) -> list[str]:
        if not text.strip():
            return []

        def split_word(word: str) -> list[str]:
            if not word:
                return []
            left, _, right, _ = draw.textbbox((0, 0), word, font=font)
            if right - left <= width:
                return [word]
            chunks: list[str] = []
            remaining = word
            while remaining:
                low = 1
                high = len(remaining)
                best = remaining[:1]
                while low <= high:
                    mid = (low + high) // 2
                    candidate = remaining[:mid]
                    c_left, _, c_right, _ = draw.textbbox((0, 0), candidate, font=font)
                    if c_right - c_left <= width:
                        best = candidate
                        low = mid + 1
                    else:
                        high = mid - 1
                chunks.append(best)
                remaining = remaining[len(best):]
            return chunks

        wrapped: list[str] = []
        current = ""
        for word in text.split():
            candidate = f"{current} {word}".strip()
            left, _, right, _ = draw.textbbox((0, 0), candidate, font=font)
            if right - left <= width:
                current = candidate
            else:
                if current:
                    wrapped.append(current)
                fragments = split_word(word)
                if len(fragments) == 1:
                    current = word
                    continue
                wrapped.extend(fragments[:-1])
                current = fragments[-1]
        if current:
            wrapped.append(current)
        return wrapped

    def _measure_lines(
        self,
        draw: ImageDraw.ImageDraw,
        lines: list[str],
        font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
        spacing: int,
    ) -> tuple[int, int]:
        if not lines:
            return 0, 0
        widths: list[int] = []
        heights: list[int] = []
        for line in lines:
            left, top, right, bottom = draw.textbbox((0, 0), line, font=font)
            widths.append(right - left)
            heights.append(bottom - top)
        total_height = sum(heights) + spacing * (len(lines) - 1)
        return max(widths), total_height

    @staticmethod
    def _spacing_candidates(font_size: int) -> list[int]:
        default_spacing = max(4, int(font_size * 0.22))
        min_spacing = max(2, int(font_size * 0.1))
        return list(range(default_spacing, min_spacing - 1, -1))

    def _truncate_lines(
        self,
        draw: ImageDraw.ImageDraw,
        lines: list[str],
        font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
        width: int,
        max_lines: int | None,
    ) -> list[str]:
        if max_lines is None or len(lines) <= max_lines:
            return lines
        trimmed = lines[:max_lines]
        if not trimmed:
            return trimmed
        overflow_text = " ".join(lines[max_lines - 1:]).strip()
        trimmed[-1] = self._truncate_text_to_width(draw, overflow_text, font, width)
        return trimmed

    @staticmethod
    def _truncate_copy_at_word_boundary(text: str, max_chars: int, *, ellipsis: bool = True) -> str:
        cleaned = " ".join(str(text or "").split()).strip()
        if len(cleaned) <= max_chars:
            return cleaned
        suffix = "..." if ellipsis else ""
        limit = max(max_chars - len(suffix), 1)
        candidate = cleaned[:limit].rstrip(" ,.;:-")
        if " " in candidate:
            candidate = candidate.rsplit(" ", 1)[0].rstrip(" ,.;:-")
        candidate = candidate or cleaned[:limit].rstrip(" ,.;:-")
        return f"{candidate}{suffix}"

    def _truncate_text_to_width(
        self,
        draw: ImageDraw.ImageDraw,
        text: str,
        font: ImageFont.FreeTypeFont | ImageFont.ImageFont,
        width: int,
    ) -> str:
        words = [word for word in str(text or "").split() if word]
        if not words:
            return "..."
        accepted: list[str] = []
        for word in words:
            candidate_words = [*accepted, word]
            candidate = " ".join(candidate_words).strip()
            candidate_with_ellipsis = f"{candidate}..."
            left, _, right, _ = draw.textbbox((0, 0), candidate_with_ellipsis, font=font)
            if right - left <= width:
                accepted = candidate_words
                continue
            break
        if accepted:
            return f"{' '.join(accepted).rstrip(' ,.;:-')}..."
        fragment = words[0]
        while fragment:
            candidate = f"{fragment.rstrip(' ,.;:-')}..."
            left, _, right, _ = draw.textbbox((0, 0), candidate, font=font)
            if right - left <= width:
                return candidate
            fragment = fragment[:-1]
        return "..."

    def _fit_text_block(
        self,
        draw: ImageDraw.ImageDraw,
        text: str,
        width: int,
        height: int,
        base_size: int,
        min_size: int = 12,
        max_lines: int | None = None,
        family_hint: str | None = None,
    ) -> tuple[ImageFont.FreeTypeFont | ImageFont.ImageFont, list[str], int, dict[str, object]]:
        best_attempt: dict[str, object] | None = None
        best_score: tuple[int, int, int] | None = None
        font_size = base_size
        while font_size >= min_size:
            font = self._font(font_size, family_hint=family_hint)
            wrapped_lines = self._wrap_text(draw, text, font, width)
            for spacing in self._spacing_candidates(font_size):
                _, total_height = self._measure_lines(draw, wrapped_lines, font, spacing)
                line_overflow = 0 if max_lines is None else max(len(wrapped_lines) - max_lines, 0)
                height_overflow = max(total_height - height, 0)
                score = (line_overflow, height_overflow, -font_size)
                if best_score is None or score < best_score:
                    best_score = score
                    best_attempt = {
                        "font": font,
                        "lines": wrapped_lines,
                        "spacing": spacing,
                        "font_size": font_size,
                        "line_count": len(wrapped_lines),
                        "truncated": False,
                        "truncation_reason": None,
                    }
                if (max_lines is None or len(wrapped_lines) <= max_lines) and total_height <= height:
                    return font, wrapped_lines, spacing, {
                        "font_size": font_size,
                        "line_count": len(wrapped_lines),
                        "truncated": False,
                        "truncation_reason": None,
                    }
            font_size -= 2

        fallback_font = self._font(min_size, family_hint=family_hint)
        fallback_spacing = self._spacing_candidates(min_size)[-1]
        wrapped_lines = self._wrap_text(draw, text, fallback_font, width)
        final_lines = list(wrapped_lines)
        truncated = False
        truncation_reason: str | None = None
        if max_lines is not None and len(final_lines) > max_lines:
            final_lines = self._truncate_lines(draw, final_lines, fallback_font, width, max_lines)
            truncated = True
            truncation_reason = "max_lines_exceeded"
        _, total_height = self._measure_lines(draw, final_lines, fallback_font, fallback_spacing)
        while total_height > height and len(final_lines) > 1:
            final_lines = final_lines[:-1]
            truncated = True
            truncation_reason = truncation_reason or "zone_height_exceeded"
            if final_lines:
                final_lines[-1] = self._truncate_text_to_width(draw, final_lines[-1], fallback_font, width)
            _, total_height = self._measure_lines(draw, final_lines, fallback_font, fallback_spacing)

        chosen_font = best_attempt.get("font") if isinstance(best_attempt, dict) else fallback_font
        chosen_spacing = int(best_attempt.get("spacing")) if isinstance(best_attempt, dict) else fallback_spacing
        chosen_lines = list(best_attempt.get("lines") or []) if isinstance(best_attempt, dict) else wrapped_lines
        chosen_font_size = int(best_attempt.get("font_size") or min_size) if isinstance(best_attempt, dict) else min_size
        chosen_line_count = int(best_attempt.get("line_count") or len(chosen_lines)) if isinstance(best_attempt, dict) else len(chosen_lines)
        best_line_overflow = 0 if max_lines is None else max(len(chosen_lines) - max_lines, 0)
        _, best_total_height = self._measure_lines(draw, chosen_lines, chosen_font, chosen_spacing)
        best_height_overflow = max(best_total_height - height, 0)
        fallback_line_overflow = 0 if max_lines is None else max(len(final_lines) - max_lines, 0)
        fallback_height_overflow = max(total_height - height, 0)
        if final_lines and (
            truncated
            or best_attempt is None
            or (fallback_line_overflow, fallback_height_overflow, -min_size) < (best_line_overflow, best_height_overflow, -chosen_font_size)
        ):
            chosen_font = fallback_font
            chosen_spacing = fallback_spacing
            chosen_lines = final_lines
            chosen_font_size = min_size
            chosen_line_count = len(final_lines)
        return chosen_font, chosen_lines, chosen_spacing, {
            "font_size": chosen_font_size,
            "line_count": chosen_line_count,
            "truncated": truncated,
            "truncation_reason": truncation_reason,
        }

    @staticmethod
    def _split_sentences(text: str) -> list[str]:
        sentences = [item.strip() for item in re.split(r"(?<=[.!?])\s+", text.strip()) if item.strip()]
        return sentences or ([text.strip()] if text.strip() else [])

    def _social_body_copy(self, body: str, preset: str) -> str:
        return " ".join(str(body or "").split()).strip()

    def _headline_copy(self, headline: str, preset: str) -> str:
        return " ".join(str(headline or "").split()).strip()

    @staticmethod
    def _draw_panel(
        draw: ImageDraw.ImageDraw,
        box: tuple[int, int, int, int],
        fill: tuple[int, int, int],
        radius: int = 28,
        outline: tuple[int, int, int] | None = None,
        width: int = 2,
    ) -> None:
        draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width if outline else 0)

    @staticmethod
    def _rounded_mask(size: tuple[int, int], radius: int) -> Image.Image:
        mask = Image.new("L", size, 0)
        ImageDraw.Draw(mask).rounded_rectangle((0, 0, size[0], size[1]), radius=radius, fill=255)
        return mask

    def _draw_brand_lockup(
        self,
        draw: ImageDraw.ImageDraw,
        brand_name: str,
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        fill: tuple[int, int, int],
    ) -> None:
        x0, y0, x1, y1 = box
        mark_size = min(y1 - y0, 44)
        mark_box = (x0, y0 + max((y1 - y0 - mark_size) // 2, 0), x0 + mark_size, y0 + max((y1 - y0 - mark_size) // 2, 0) + mark_size)
        self._draw_panel(draw, mark_box, primary, radius=12)
        inner = 8
        draw.rounded_rectangle(
            (mark_box[0] + inner, mark_box[1] + inner, mark_box[2] - inner, mark_box[3] - inner),
            radius=8,
            fill=accent,
        )
        text_zone = type(
            "Zone",
            (),
            {
                "x": mark_box[2] + 14,
                "y": y0,
                "width": max(x1 - (mark_box[2] + 14), 10),
                "height": y1 - y0,
                "max_lines": 1,
            },
        )()
        self._draw_text_block(draw, brand_name.upper(), text_zone, fill, 30, padding=0)

    def _draw_badges(
        self,
        draw: ImageDraw.ImageDraw,
        badges: list[str],
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
    ) -> None:
        if not badges:
            return
        x0, y0, x1, y1 = box
        cursor_x = x0
        cursor_y = y0
        max_width = x1 - x0
        font = self._font(18)
        for badge in badges:
            left, top, right, bottom = draw.textbbox((0, 0), badge, font=font)
            text_width = right - left
            text_height = bottom - top
            pill_width = text_width + 28
            pill_height = text_height + 18
            if cursor_x + pill_width > x0 + max_width:
                break
            fill = self._blend_color((255, 255, 255), accent, 0.18)
            outline = self._blend_color(fill, primary, 0.25)
            draw.rounded_rectangle((cursor_x, cursor_y, cursor_x + pill_width, cursor_y + pill_height), radius=16, fill=fill, outline=outline, width=1)
            draw.text((cursor_x + 14, cursor_y + 9 - top), badge, fill=primary, font=font)
            cursor_x += pill_width + 12

    def _draw_section_label(
        self,
        draw: ImageDraw.ImageDraw,
        label: str,
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
    ) -> None:
        if not label:
            return
        x0, y0, _, y1 = box
        font = self._font(18)
        left, top, right, bottom = draw.textbbox((0, 0), label.upper(), font=font)
        pill_width = (right - left) + 28
        pill_height = max((bottom - top) + 16, y1 - y0)
        pill_box = (x0, y0, x0 + pill_width, y0 + pill_height)
        fill = self._blend_color((255, 255, 255), accent, 0.2)
        outline = self._blend_color(fill, primary, 0.18)
        draw.rounded_rectangle(pill_box, radius=18, fill=fill, outline=outline, width=1)
        draw.text((pill_box[0] + 14, pill_box[1] + 8 - top), label.upper(), fill=primary, font=font)

    def _draw_bullet_list(
        self,
        draw: ImageDraw.ImageDraw,
        items: list[str],
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        base_size: int = 20,
        family_hint: str | None = None,
        badge_style: dict | None = None,
    ) -> list[dict[str, int | str | None]]:
        if not items:
            return []
        font = self._font(base_size, family_hint=family_hint)

        # Fix 4: Support numbered badges from component motifs
        use_numbered_badges = badge_style and badge_style.get("shape") == "rounded_rect"
        # TODO: Load default_badge_color from brand_visual_brief instead of hardcoded #F7941D
        default_badge = accent or "#F7941D"  # Fallback orange
        badge_color = self._parse_color(badge_style.get("badge_color", default_badge), accent) if badge_style else accent
        badge_text_color = self._parse_color(badge_style.get("text_color", "#FFFFFF"), (255, 255, 255)) if badge_style else (255, 255, 255)
        badge_radius = badge_style.get("radius_px", 12) if badge_style else 12
        badge_padding = badge_style.get("padding_px", 8) if badge_style else 8

        bullet_radius = 5
        x0, y0, x1, y1 = box
        badge_offset = 40 if use_numbered_badges else 18
        available_width = max(x1 - x0 - badge_offset - 6, 10)
        cursor_y = y0
        zones: list[dict[str, int | str | None]] = []

        for index, item in enumerate(items, start=1):
            wrapped = self._wrap_text(draw, item, font, available_width)
            if not wrapped:
                continue
            _, total_height = self._measure_lines(draw, wrapped, font, 6)
            if cursor_y + total_height > y1:
                break
            first_line = wrapped[0]
            left, top, _, bottom = draw.textbbox((0, 0), first_line, font=font)
            first_line_height = max(bottom - top, bullet_radius * 2)

            if use_numbered_badges:
                # Draw numbered badge
                badge_font = self._font(max(14, base_size - 4), family_hint=family_hint, weight="bold")
                number_text = f"{index:02d}" if badge_style.get("number_format") == "01" else str(index)

                # Measure badge text
                badge_left, badge_top, badge_right, badge_bottom = draw.textbbox((0, 0), number_text, font=badge_font)
                badge_text_width = badge_right - badge_left
                badge_text_height = badge_bottom - badge_top

                # Calculate badge box
                badge_width = badge_text_width + (badge_padding * 2)
                badge_height = badge_text_height + (badge_padding * 2)
                badge_y = cursor_y + max(0, int((first_line_height - badge_height) / 2))

                # Draw rounded rectangle badge
                badge_box = (x0, badge_y, x0 + badge_width, badge_y + badge_height)
                draw.rounded_rectangle(badge_box, radius=badge_radius, fill=badge_color)

                # Draw number text in badge
                text_x = x0 + badge_padding - badge_left
                text_y = badge_y + badge_padding - badge_top
                draw.text((text_x, text_y), number_text, fill=badge_text_color, font=badge_font)

                line_x = x0 + badge_offset
            else:
                # Draw traditional bullet
                bullet_y = cursor_y + max(0, int((first_line_height - (bullet_radius * 2)) / 2))
                draw.ellipse((x0, bullet_y, x0 + bullet_radius * 2, bullet_y + bullet_radius * 2), fill=accent)
                line_x = x0 + 18

            # Draw list item text
            line_cursor = cursor_y
            for line in wrapped:
                left, top, _, bottom = draw.textbbox((0, 0), line, font=font)
                draw.text((line_x, line_cursor - top), line, fill=primary, font=font)
                line_cursor += (bottom - top) + 6

            zones.append(self._zone_manifest(f"proof_point_{index}", "proof_point", (x0, cursor_y, x1, cursor_y + total_height), len(wrapped)))
            cursor_y += total_height + 12

        return zones

    def _draw_stat_cards(
        self,
        draw: ImageDraw.ImageDraw,
        stats: list[str],
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        background: tuple[int, int, int],
    ) -> list[dict[str, int | str | None]]:
        if not stats:
            return []
        x0, y0, x1, y1 = box
        card_gap = 12
        count = max(1, min(len(stats), 3))
        total_gap = card_gap * (count - 1)
        card_width = max((x1 - x0 - total_gap) // count, 80)
        card_height = max(y1 - y0, 56)
        zones: list[dict[str, int | str | None]] = []
        for index, stat in enumerate(stats[:count], start=1):
            card_x0 = x0 + (index - 1) * (card_width + card_gap)
            card_box = (card_x0, y0, card_x0 + card_width, y0 + card_height)
            fill = self._blend_color(background, accent, 0.16 if index % 2 else 0.09)
            outline = self._blend_color(fill, primary, 0.16)
            draw.rounded_rectangle(card_box, radius=18, fill=fill, outline=outline, width=1)
            zone = type(
                "Zone",
                (),
                self._zone_manifest(f"stat_{index}", "stat_highlight", (card_box[0] + 12, card_box[1] + 10, card_box[2] - 12, card_box[3] - 10), 2),
            )()
            self._draw_text_block(draw, stat, zone, primary, 18, padding=0)
            zones.append(self._zone_manifest(f"stat_{index}", "stat_highlight", card_box, 2))
        return zones

    def _draw_editorial_motif(
        self,
        draw: ImageDraw.ImageDraw,
        box: tuple[int, int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        background: tuple[int, int, int],
    ) -> None:
        x0, y0, x1, y1 = box
        width = x1 - x0
        height = y1 - y0
        soft_primary = self._blend_color(background, primary, 0.16)
        soft_accent = self._blend_color(background, accent, 0.22)
        draw.rounded_rectangle(box, radius=32, fill=self._blend_color(background, primary, 0.06))
        draw.ellipse((x0 + width * 0.08, y0 + height * 0.1, x0 + width * 0.5, y0 + height * 0.72), fill=soft_accent)
        draw.rounded_rectangle((x0 + width * 0.62, y0 + height * 0.18, x0 + width * 0.84, y0 + height * 0.82), radius=26, fill=soft_primary)
        arrow_points = [
            (int(x0 + width * 0.16), int(y0 + height * 0.68)),
            (int(x0 + width * 0.38), int(y0 + height * 0.5)),
            (int(x0 + width * 0.54), int(y0 + height * 0.58)),
            (int(x0 + width * 0.8), int(y0 + height * 0.3)),
        ]
        draw.line(arrow_points, fill=primary, width=max(8, width // 45), joint="curve")
        tip = arrow_points[-1]
        draw.polygon(
            [
                tip,
                (tip[0] - max(20, width // 18), tip[1] + max(10, height // 30)),
                (tip[0] - max(10, width // 26), tip[1] - max(24, height // 18)),
            ],
            fill=primary,
        )
        bar_base = y1 - max(28, height // 10)
        bar_width = max(26, width // 16)
        for index, factor in enumerate((0.18, 0.34, 0.52), start=0):
            bx0 = int(x0 + width * 0.12) + index * (bar_width + 14)
            bx1 = bx0 + bar_width
            by0 = int(bar_base - (height * factor))
            draw.rounded_rectangle((bx0, by0, bx1, bar_base), radius=10, fill=primary if index == 2 else soft_primary)
        draw.arc(
            (x0 + width * 0.46, y0 + height * 0.2, x0 + width * 0.94, y0 + height * 0.92),
            start=210,
            end=338,
            fill=self._blend_color(background, accent, 0.44),
            width=max(4, width // 90),
        )
        draw.ellipse(
            (x0 + width * 0.72, y0 + height * 0.12, x0 + width * 0.79, y0 + height * 0.19),
            fill=self._blend_color(background, accent, 0.72),
        )
        draw.ellipse(
            (x0 + width * 0.66, y0 + height * 0.72, x0 + width * 0.72, y0 + height * 0.78),
            fill=self._blend_color(background, primary, 0.42),
        )

    def _paste_visual_card(
        self,
        canvas: Image.Image,
        storage_path: str,
        box: tuple[int, int, int, int],
        radius: int = 28,
        frame_fill: tuple[int, int, int] | None = None,
        padding: int = 0,
    ) -> dict[str, object]:
        source = Path(self.storage.absolute_path(storage_path))
        if not source.exists():
            return {"rendered": False, "storage_path": storage_path, "reason": "missing"}
        x0, y0, x1, y1 = box
        if frame_fill:
            ImageDraw.Draw(canvas).rounded_rectangle(box, radius=radius, fill=frame_fill)
        inner_box = (x0 + padding, y0 + padding, x1 - padding, y1 - padding)
        inner_width = max(inner_box[2] - inner_box[0], 1)
        inner_height = max(inner_box[3] - inner_box[1], 1)
        with open_image_asset(source) as raw:
            visual = raw.convert("RGB")
            fitted, assessment = self._compose_visual_for_zone(
                visual,
                inner_width,
                inner_height,
                radius=max(radius - padding, 12),
                frame_fill=frame_fill,
            )
        mask = self._rounded_mask((inner_width, inner_height), max(radius - padding, 12))
        canvas.paste(fitted, (inner_box[0], inner_box[1]), mask)
        return {
            "rendered": True,
            "storage_path": storage_path,
            **assessment,
        }

    def _paginate_body(self, body: str, max_chars: int, max_pages: int) -> list[str]:
        sentences = self._split_sentences(body)
        if not sentences:
            return [""]
        pages: list[str] = []
        current = ""
        for sentence in sentences:
            candidate = f"{current} {sentence}".strip()
            if current and len(candidate) > max_chars:
                pages.append(current)
                current = sentence
            else:
                current = candidate
        if current:
            pages.append(current)
        if len(pages) <= max_pages:
            return pages
        collapsed = pages[: max_pages - 1]
        collapsed.append(" ".join(pages[max_pages - 1 :]))
        return collapsed

    @staticmethod
    def _collapse_page_segments(segments: list[str], max_pages: int) -> list[str]:
        cleaned: list[str] = []
        seen: set[str] = set()
        for segment in segments:
            normalized = " ".join(str(segment or "").split()).strip()
            if not normalized:
                continue
            key = normalized.lower()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(normalized)
        if not cleaned:
            return [""]
        if len(cleaned) <= max_pages:
            return cleaned
        collapsed = cleaned[: max_pages - 1]
        collapsed.append(" ".join(cleaned[max_pages - 1 :]).strip())
        return collapsed

    @staticmethod
    def _structured_render_sections(payload: RendererInput) -> dict[str, object]:
        metadata = payload.text.metadata if isinstance(payload.text.metadata, dict) else {}
        render_sections = metadata.get("render_sections", {})
        return render_sections if isinstance(render_sections, dict) else {}

    @staticmethod
    def _page_list_value(value: object) -> list[str]:
        if not isinstance(value, list):
            return []
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in value:
            normalized = " ".join(str(item or "").split()).strip()
            if not normalized:
                continue
            key = normalized.casefold()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(normalized)
        return cleaned

    def _build_carousel_segments(self, payload: RendererInput, max_pages: int) -> list[str]:
        render_sections = self._structured_render_sections(payload)
        carousel_slide_specs = payload.text.metadata.get("carousel_slide_specs", []) if isinstance(payload.text.metadata, dict) else []
        if isinstance(carousel_slide_specs, list) and carousel_slide_specs:
            segments = [
                " ".join(
                    part
                    for part in [
                        str(slide.get("headline") or "").strip(),
                        self._carousel_slide_body(
                            slide,
                            render_sections=render_sections,
                            fallback_body=payload.text.body,
                        ),
                        " ".join(str(item).strip() for item in (slide.get("proof_points") or []) if str(item).strip()).strip(),
                        str(slide.get("cta") or "").strip(),
                    ]
                    if part
                ).strip()
                for slide in carousel_slide_specs
                if isinstance(slide, dict)
            ]
            return self._collapse_page_segments(segments, max_pages=max_pages)
        body = str(payload.text.body or "").strip()
        sentences = self._split_sentences(body)
        supporting_line = str(render_sections.get("supporting_line") or "").strip() or self._supporting_line(payload, "carousel", body)
        proof_points = self._page_list_value(render_sections.get("proof_points")) or self._proof_points(payload, "carousel", body)
        stat_highlights = self._page_list_value(render_sections.get("stat_highlights")) or self._stat_highlights(payload, "carousel")

        lead_segment = supporting_line or (sentences[0] if sentences else body)
        remaining_sentences = sentences[1:] if len(sentences) > 1 else []

        detail_segments: list[str] = []
        for index in range(0, len(remaining_sentences), 2):
            segment = " ".join(remaining_sentences[index : index + 2]).strip()
            if segment:
                detail_segments.append(segment)

        if not detail_segments and proof_points:
            for index in range(0, len(proof_points), 2):
                segment = ". ".join(point.strip().rstrip(".") for point in proof_points[index : index + 2] if point).strip()
                if segment:
                    if not segment.endswith("."):
                        segment = f"{segment}."
                    detail_segments.append(segment)

        if not detail_segments and body:
            detail_segments.append(body)

        closing_segment = ""
        if stat_highlights:
            closing_segment = " | ".join(item.strip() for item in stat_highlights[:3] if item).strip()
        elif payload.text.cta:
            closing_segment = str(payload.text.cta).strip()
        elif proof_points:
            closing_segment = " | ".join(point.strip() for point in proof_points[:2] if point).strip()

        segments = self._collapse_page_segments(
            [
                lead_segment,
                *detail_segments,
                closing_segment,
            ],
            max_pages=max_pages,
        )

        if len(segments) < min(3, max_pages):
            fallback_segments = self._collapse_page_segments(
                [
                    lead_segment,
                    body,
                    " ".join(point.strip() for point in proof_points[:2] if point).strip(),
                    " | ".join(item.strip() for item in stat_highlights[:2] if item).strip(),
                    str(payload.text.cta or "").strip(),
                ],
                max_pages=max_pages,
            )
            for fallback in fallback_segments:
                if len(segments) >= min(3, max_pages):
                    break
                key = fallback.lower()
                if not fallback or any(existing.lower() == key for existing in segments):
                    continue
                segments.append(fallback)

        return self._collapse_page_segments(segments, max_pages=max_pages)

    def _carousel_slide_body(
        self,
        slide: dict,
        *,
        render_sections: dict[str, object],
        fallback_body: str,
    ) -> str:
        for key in ("body", "body_display", "detail", "narrative", "summary", "description"):
            value = slide.get(key)
            if isinstance(value, list):
                joined = ". ".join(self._page_list_value(value)).strip()
                if joined:
                    return joined
            text = str(value or "").strip()
            if text:
                return text
        body_points = self._page_list_value(slide.get("body_points")) or self._page_list_value(slide.get("detail_points"))
        if body_points:
            return ". ".join(body_points).strip()
        proof_points = self._page_list_value(slide.get("proof_points"))
        if proof_points:
            return ". ".join(proof_points).strip()
        supporting_line = str(slide.get("supporting_line") or slide.get("subheadline") or "").strip()
        if supporting_line:
            return supporting_line
        return str(render_sections.get("body_display") or fallback_body or "").strip()

    def _carousel_slide_supporting_line(
        self,
        slide: dict,
        *,
        body_text: str,
    ) -> str:
        supporting_line = str(slide.get("supporting_line") or slide.get("subheadline") or "").strip()
        if supporting_line and supporting_line.casefold() != body_text.casefold():
            return supporting_line
        return ""

    def _build_page_payloads(self, payload: RendererInput) -> list[dict]:
        format_name = payload.studio_panel.get("format", "static")
        metadata = payload.text.metadata if isinstance(payload.text.metadata, dict) else {}
        composition = self._composition_plan(payload)
        render_sections = self._structured_render_sections(payload)
        primary_visual_plan = composition.get("primary_visual_plan", {}) if isinstance(composition, dict) else {}
        text_content_plan = composition.get("text_content_plan", {}) if isinstance(composition, dict) else {}
        show_primary_visual = bool(text_content_plan.get("show_primary_visual", True))
        show_primary_visual_on_first_page_only = bool(primary_visual_plan.get("show_primary_visual_on_first_page_only", True))
        if format_name == "infographic":
            return [
                {
                    "headline": render_sections.get("headline_display") or payload.text.headline,
                    "body": render_sections.get("body_display") or payload.text.body,
                    "supporting_line": render_sections.get("supporting_line"),
                    "proof_points": self._page_list_value(render_sections.get("proof_points")),
                    "stat_highlights": self._page_list_value(render_sections.get("stat_highlights")),
                    "infographic_section_specs": metadata.get("infographic_section_specs", []) if isinstance(metadata.get("infographic_section_specs"), list) else [],
                    "cta": render_sections.get("cta_display") or payload.text.cta,
                    "show_image": show_primary_visual,
                    "page_index": 0,
                    "page_count": 1,
                    "content_structure_type": render_sections.get("creative_type") or format_name,
                }
            ]
        if format_name not in {"carousel", "pdf"}:
            return [
                {
                    "headline": render_sections.get("headline_display") or payload.text.headline,
                    "body": render_sections.get("body_display") or payload.text.body,
                    "supporting_line": render_sections.get("supporting_line"),
                    "proof_points": self._page_list_value(render_sections.get("proof_points")),
                    "stat_highlights": self._page_list_value(render_sections.get("stat_highlights")),
                    "static_panel_spec": metadata.get("static_panel_spec", {}) if isinstance(metadata.get("static_panel_spec"), dict) else {},
                    "cta": render_sections.get("cta_display") or payload.text.cta,
                    "show_image": show_primary_visual,
                    "content_structure_type": render_sections.get("creative_type") or format_name,
                }
            ]
        max_pages = 6 if format_name == "carousel" else 4
        carousel_slide_specs = payload.text.metadata.get("carousel_slide_specs", []) if isinstance(payload.text.metadata, dict) else []
        if format_name == "carousel" and isinstance(carousel_slide_specs, list) and carousel_slide_specs:
            pages: list[dict] = []
            total_segments = max(len(carousel_slide_specs), 1)
            for index, slide in enumerate(carousel_slide_specs):
                if not isinstance(slide, dict):
                    continue
                body_text = self._carousel_slide_body(
                    slide,
                    render_sections=render_sections,
                    fallback_body=payload.text.body,
                )
                pages.append(
                    {
                        "headline": str(slide.get("headline") or payload.text.headline).strip(),
                        "body": body_text,
                        "supporting_line": self._carousel_slide_supporting_line(slide, body_text=body_text),
                        "proof_points": self._page_list_value(slide.get("proof_points")),
                        "stat_highlights": self._page_list_value(slide.get("stat_highlights")),
                        "cta": str(slide.get("cta") or "").strip(),
                        "show_image": show_primary_visual and (index == 0 or not show_primary_visual_on_first_page_only),
                        "page_index": index,
                        "page_count": total_segments,
                        "content_role": str(slide.get("content_role") or slide.get("role") or "detail").strip(),
                        "content_structure_type": "carousel",
                    }
                )
            return pages or [
                {
                    "headline": payload.text.headline,
                    "body": payload.text.body,
                    "cta": payload.text.cta,
                    "show_image": show_primary_visual,
                    "content_structure_type": "carousel",
                }
            ]
        if format_name == "carousel":
            segments = self._build_carousel_segments(payload, max_pages=max_pages)
        else:
            segments = self._paginate_body(payload.text.body, max_chars=320, max_pages=max_pages)
        pages: list[dict] = []
        total_segments = max(len(segments), 1)
        for index, segment in enumerate(segments):
            is_first = index == 0
            is_last = index == total_segments - 1
            pages.append(
                {
                    "headline": payload.text.headline if is_first else f"{payload.text.headline} Continued",
                    "body": segment,
                    "supporting_line": render_sections.get("supporting_line") if is_first else "",
                    "proof_points": self._page_list_value(render_sections.get("proof_points")) if is_first else [],
                    "stat_highlights": self._page_list_value(render_sections.get("stat_highlights")) if is_last else [],
                    "cta": payload.text.cta if is_last else "",
                    "show_image": show_primary_visual and (is_first or not show_primary_visual_on_first_page_only),
                    "page_index": index,
                    "page_count": total_segments,
                    "content_structure_type": format_name,
                }
            )
        return pages

    def _draw_text_block(
        self,
        draw: ImageDraw.ImageDraw,
        text: str,
        zone,
        fill: tuple[int, int, int],
        base_size: int,
        align: str = "left",
        background_fill: tuple[int, int, int] | None = None,
        padding: int = 18,
        style_hint: dict | None = None,
        font_family: str | None = None,
    ) -> dict[str, object]:
        family_hint = str(font_family or "").strip() or None
        # Fix 5: Support configurable background_radius from style
        background_radius = 18  # Default radius
        if style_hint:
            fill = self._resolved_fill_from_style_hint(style_hint, fill)
            if background_fill is None:
                background_fill = self._parse_color(
                    self._style_color_payload_to_hex(style_hint.get("background_color")),
                    background_fill or fill,
                ) if style_hint.get("background_color") else background_fill
            estimated_font_size = int(style_hint.get("estimated_font_size") or 0)
            if estimated_font_size:
                base_size = max(base_size, estimated_font_size)
            if not family_hint:
                family_hint = str(style_hint.get("font_family") or "").strip() or None
            # Get custom background radius from style
            if style_hint.get("background_radius"):
                try:
                    background_radius = int(style_hint.get("background_radius"))
                except (ValueError, TypeError):
                    pass
        if background_fill:
            draw.rounded_rectangle(self._zone_box(zone), radius=background_radius, fill=background_fill)
        inner_x = zone.x + padding
        inner_y = zone.y + padding
        inner_w = max(zone.width - (padding * 2), 10)
        inner_h = max(zone.height - (padding * 2), 10)
        # Format-aware minimum font size - allow smaller fonts for data-rich formats
        active_payload = self.payload if isinstance(self.payload, RendererInput) else None
        studio_panel = active_payload.studio_panel if active_payload else {}
        format_name = str(studio_panel.get("format") or "").strip().lower()
        role_name = str(getattr(zone, "role", "") or "").strip().lower()
        if role_name in {"legal", "footer", "disclaimer"}:
            min_size_for_format = 7
        else:
            min_size_for_format = 9 if format_name in ["infographic", "static"] else 12
        font, lines, spacing, fit_meta = self._fit_text_block(
            draw=draw,
            text=text,
            width=inner_w,
            height=inner_h,
            base_size=base_size,
            min_size=min_size_for_format,
            max_lines=zone.max_lines,
            family_hint=family_hint,
        )
        _, total_height = self._measure_lines(draw, lines, font, spacing)
        cursor_y = inner_y + max((inner_h - total_height) // 2, 0)
        occupied_bounds: list[tuple[int, int, int, int]] = []
        for line in lines:
            left, top, right, bottom = draw.textbbox((0, 0), line, font=font)
            line_width = right - left
            line_height = bottom - top
            cursor_x = inner_x
            if align == "center":
                cursor_x = inner_x + max((inner_w - line_width) // 2, 0)
            elif align == "right":
                cursor_x = inner_x + max(inner_w - line_width, 0)
            text_top = cursor_y - top
            draw.text((cursor_x, text_top), line, fill=fill, font=font)
            occupied_bounds.append((cursor_x, text_top, cursor_x + line_width, text_top + line_height))
            cursor_y += line_height + spacing
        occupied_box = None
        if occupied_bounds:
            occupied_box = (
                min(box[0] for box in occupied_bounds),
                min(box[1] for box in occupied_bounds),
                max(box[2] for box in occupied_bounds),
                max(box[3] for box in occupied_bounds),
            )
        fitted_text = " ".join(lines).strip()
        manifest = {
            "role": getattr(zone, "role", ""),
            "zone_id": getattr(zone, "zone_id", ""),
            "original_text_length": len(str(text or "")),
            "fitted_text_length": len(fitted_text),
            "font_size": fit_meta.get("font_size"),
            "line_count": fit_meta.get("line_count"),
            "line_spacing": spacing,
            "truncated": bool(fit_meta.get("truncated")),
            "truncation_reason": fit_meta.get("truncation_reason"),
            "occupied_box": occupied_box,
            "zone_box": self._zone_box(zone),
        }
        logger.debug(
            "renderer.text_fit role=%s zone_id=%s original_length=%s fitted_length=%s font_size=%s line_count=%s spacing=%s truncated=%s truncation_reason=%s occupied_box=%s",
            manifest["role"],
            manifest["zone_id"],
            manifest["original_text_length"],
            manifest["fitted_text_length"],
            manifest["font_size"],
            manifest["line_count"],
            manifest["line_spacing"],
            manifest["truncated"],
            manifest["truncation_reason"],
            manifest["occupied_box"],
        )
        return manifest

    def _base_canvas_text_cleanup_fill(
        self,
        *,
        fill: tuple[int, int, int],
        background: tuple[int, int, int],
        primary: tuple[int, int, int],
        role: str,
    ) -> tuple[tuple[int, int, int], tuple[int, int, int], int]:
        if self._relative_luminance(fill) >= 0.62:
            cleanup_fill = self._blend_color(background, primary, 0.84)
            resolved_fill = self._ensure_text_contrast(fill, cleanup_fill, fallback=(255, 255, 255))
            cleanup_padding = 4 if role in {"legal", "footer", "disclaimer"} else 16 if role == "headline" else 12
            return cleanup_fill, resolved_fill, cleanup_padding
        cleanup_fill = self._blend_color(background, (255, 255, 255), 0.72)
        resolved_fill = self._ensure_text_contrast(fill, cleanup_fill, fallback=(20, 20, 20))
        cleanup_padding = 4 if role in {"legal", "footer", "disclaimer"} else 16 if role == "headline" else 12
        return cleanup_fill, resolved_fill, cleanup_padding

    @staticmethod
    def _boxes_overlap(
        first: tuple[int, int, int, int] | None,
        second: tuple[int, int, int, int] | None,
        *,
        padding: int = 0,
    ) -> bool:
        if not first or not second:
            return False
        left_a, top_a, right_a, bottom_a = first
        left_b, top_b, right_b, bottom_b = second
        return not (
            (right_a + padding) <= left_b
            or (right_b + padding) <= left_a
            or (bottom_a + padding) <= top_b
            or (bottom_b + padding) <= top_a
        )

    @classmethod
    def _box_overlaps_any(
        cls,
        box: tuple[int, int, int, int] | None,
        reserved_boxes: list[tuple[int, int, int, int]],
        *,
        padding: int = 0,
    ) -> bool:
        return any(cls._boxes_overlap(box, reserved, padding=padding) for reserved in reserved_boxes if reserved)

    @staticmethod
    def _clamp_box(
        box: tuple[int, int, int, int],
        *,
        canvas_width: int,
        canvas_height: int,
    ) -> tuple[int, int, int, int]:
        left, top, right, bottom = box
        width = max(right - left, 1)
        height = max(bottom - top, 1)
        left = min(max(left, 0), max(canvas_width - width, 0))
        top = min(max(top, 0), max(canvas_height - height, 0))
        return (left, top, left + width, top + height)

    @classmethod
    def _adapt_box_away_from_reserved(
        cls,
        box: tuple[int, int, int, int],
        *,
        reserved_boxes: list[tuple[int, int, int, int]],
        canvas_width: int,
        canvas_height: int,
        padding: int = 18,
    ) -> tuple[tuple[int, int, int, int], bool]:
        clamped = cls._clamp_box(box, canvas_width=canvas_width, canvas_height=canvas_height)
        if not cls._box_overlaps_any(clamped, reserved_boxes, padding=padding):
            return clamped, False
        left, top, right, bottom = clamped
        width = max(right - left, 1)
        height = max(bottom - top, 1)
        for scale in (0.92, 0.84, 0.76, 0.68):
            scaled_width = max(int(width * scale), 48)
            scaled_height = max(int(height * scale), 36)
            candidate_positions = [
                (left, top),
                (max(canvas_width - scaled_width - padding, 0), top),
                (left, 0),
                (max(canvas_width - scaled_width - padding, 0), 0),
                (left, max(canvas_height - scaled_height - padding, 0)),
                (max(canvas_width - scaled_width - padding, 0), max(canvas_height - scaled_height - padding, 0)),
            ]
            for candidate_left, candidate_top in candidate_positions:
                candidate = (
                    int(candidate_left),
                    int(candidate_top),
                    int(candidate_left + scaled_width),
                    int(candidate_top + scaled_height),
                )
                candidate = cls._clamp_box(candidate, canvas_width=canvas_width, canvas_height=canvas_height)
                if not cls._box_overlaps_any(candidate, reserved_boxes, padding=padding):
                    return candidate, True
        return clamped, True

    @staticmethod
    def _zone_manifest(
        zone_id: str,
        role: str,
        box: tuple[int, int, int, int],
        max_lines: int | None = None,
    ) -> dict[str, int | str | None]:
        x0, y0, x1, y1 = box
        return {
            "zone_id": zone_id,
            "role": role,
            "x": x0,
            "y": y0,
            "width": x1 - x0,
            "height": y1 - y0,
            "max_lines": max_lines,
        }

    @staticmethod
    def _fit_image_to_zone(source: Image.Image, width: int, height: int) -> Image.Image:
        return ImageOps.fit(source, (width, height), method=Image.Resampling.LANCZOS)

    @staticmethod
    def _assess_visual_fit(source_width: int, source_height: int, target_width: int, target_height: int) -> dict[str, float | str]:
        source_width = max(int(source_width), 1)
        source_height = max(int(source_height), 1)
        target_width = max(int(target_width), 1)
        target_height = max(int(target_height), 1)
        cover_scale = max(target_width / source_width, target_height / source_height)
        visible_width = min(source_width, target_width / cover_scale)
        visible_height = min(source_height, target_height / cover_scale)
        crop_retention = max(
            0.0,
            min(1.0, (visible_width * visible_height) / float(source_width * source_height)),
        )
        resolution_ratio = max(
            0.0,
            min(1.0, min(source_width / target_width, source_height / target_height)),
        )
        aspect_ratio_delta = abs((source_width / source_height) - (target_width / target_height))
        score = max(
            0.0,
            min(1.0, (crop_retention * 0.72) + (resolution_ratio * 0.28) - min(aspect_ratio_delta * 0.08, 0.22)),
        )
        fit_mode = "cover" if crop_retention >= 0.68 and score >= 0.58 else "contain"
        return {
            "fit_mode": fit_mode,
            "score": round(score, 3),
            "crop_retention": round(crop_retention, 3),
            "resolution_ratio": round(resolution_ratio, 3),
        }

    def _compose_visual_for_zone(
        self,
        source: Image.Image,
        width: int,
        height: int,
        *,
        radius: int = 0,
        frame_fill: tuple[int, int, int] | None = None,
    ) -> tuple[Image.Image, dict[str, float | str]]:
        assessment = self._assess_visual_fit(source.width, source.height, width, height)
        fit_mode = str(assessment["fit_mode"])
        if fit_mode == "cover":
            composed = self._fit_image_to_zone(source, width, height)
        else:
            panel_color = frame_fill or (245, 241, 233)
            composed = Image.new("RGB", (width, height), color=panel_color)
            inset = max(18, min(width, height) // 18)
            inner_width = max(width - (inset * 2), 1)
            inner_height = max(height - (inset * 2), 1)
            contained = ImageOps.contain(source, (inner_width, inner_height), method=Image.Resampling.LANCZOS)
            offset_x = max((width - contained.width) // 2, 0)
            offset_y = max((height - contained.height) // 2, 0)
            composed.paste(contained, (offset_x, offset_y))
        if radius > 0:
            mask = self._rounded_mask((width, height), radius)
            masked = Image.new("RGB", (width, height), color=frame_fill or (255, 255, 255))
            masked.paste(composed, (0, 0), mask)
            composed = masked
        return composed, assessment

    def _paste_visual(
        self,
        canvas: Image.Image,
        storage_path: str,
        zone,
        *,
        frame_fill: tuple[int, int, int] | None = None,
        radius: int = 0,
    ) -> dict[str, object]:
        source = Path(self.storage.absolute_path(storage_path))
        if not source.exists():
            return {"rendered": False, "storage_path": storage_path, "reason": "missing"}
        with open_image_asset(source) as raw:
            visual = raw.convert("RGB")
            fitted, assessment = self._compose_visual_for_zone(
                visual,
                zone.width,
                zone.height,
                radius=radius,
                frame_fill=frame_fill,
            )
            canvas.paste(fitted, (zone.x, zone.y))
        return {
            "rendered": True,
            "storage_path": storage_path,
            **assessment,
        }

    def _paste_template_background(self, canvas: Image.Image, template_asset_path: str | None) -> bool:
        if not template_asset_path:
            return False
        source = self._resolve_template_background_source(template_asset_path)
        if not source:
            return False
        try:
            with open_image_asset(source) as raw:
                background = raw.convert("RGB")
                fitted = self._fit_image_to_zone(background, canvas.width, canvas.height)
                canvas.paste(fitted, (0, 0))
        except OSError:
            return False
        return True

    def _resolve_template_background_source(self, template_asset_path: str) -> Path | None:
        source = Path(self.storage.absolute_path(template_asset_path))
        if source.exists() and source.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
            return source
        scratch_root = source.parent / "_ocr"
        page_images_dir = scratch_root / "page_images"
        if page_images_dir.exists():
            first_page = next((path for path in sorted(page_images_dir.glob("page_*.png")) if path.exists()), None)
            if first_page:
                return first_page
        if scratch_root.exists():
            first_image = next(
                (
                    path for path in sorted(scratch_root.glob("*.png"))
                    if path.exists()
                ),
                None,
            )
            if first_image:
                return first_image
        return source if source.exists() and source.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"} else None

    def _paste_logo(self, canvas: Image.Image, logo_path: str | None, zone) -> bool:
        if not logo_path:
            return False
        source = Path(self.storage.absolute_path(logo_path))
        if not source.exists():
            logger.warning("renderer.logo.missing storage_path=%s", logo_path)
            return False
        try:
            with open_image_asset(source) as raw:
                logo = self._trim_transparent_logo_margins(
                    self._strip_logo_background_if_safe(raw.convert("RGBA"))
                )
                contained = ImageOps.contain(logo, (zone.width, zone.height), method=Image.Resampling.LANCZOS)
                offset_x = zone.x + max((zone.width - contained.width) // 2, 0)
                offset_y = zone.y + max((zone.height - contained.height) // 2, 0)
                canvas.paste(contained, (offset_x, offset_y), contained)
        except OSError:
            logger.warning("renderer.logo.unreadable storage_path=%s", logo_path)
            return False
        return True

    def _paste_decorative_asset(
        self,
        canvas: Image.Image,
        asset: GeneratedImageAsset,
        box: tuple[int, int, int, int],
        opacity: int = 176,
    ) -> bool:
        source = Path(self.storage.absolute_path(asset.storage_path))
        if not source.exists():
            return False
        x0, y0, x1, y1 = box
        width = max(x1 - x0, 1)
        height = max(y1 - y0, 1)
        with open_image_asset(source) as raw:
            decorative = raw.convert("RGBA")
            contained = ImageOps.contain(decorative, (width, height), method=Image.Resampling.LANCZOS)
            alpha = contained.getchannel("A").point(lambda value: min(value, opacity))
            contained.putalpha(alpha)
            offset_x = x0 + max((width - contained.width) // 2, 0)
            offset_y = y0 + max((height - contained.height) // 2, 0)
            canvas.paste(contained, (offset_x, offset_y), contained)
        return True

    def _apply_decorative_assets(
        self,
        canvas: Image.Image,
        payload: RendererInput,
        *,
        reserved_boxes: list[tuple[int, int, int, int]] | None = None,
    ) -> list[dict[str, int | str | None]]:
        if not payload.decorative_assets:
            return []
        composition = self._composition_plan(payload)
        decorative_plan = composition.get("decorative_plan", {}) if isinstance(composition, dict) else {}
        policy = str(decorative_plan.get("policy") or "")
        if policy in {"template_only", "none"}:
            return []
        source_mode = payload.blueprint.source_mode
        if source_mode == "exact_template":
            return []
        width, height = canvas.size
        placements = [
            (width - 240, 26, width - 36, 190),
            (26, height - 210, 236, height - 36),
            (width - 220, height - 220, width - 36, height - 36),
        ]
        zones: list[dict[str, int | str | None]] = []
        max_assets = int(decorative_plan.get("max_assets") or len(placements))
        max_assets = max(0, min(max_assets, len(placements)))
        for index, asset in enumerate(payload.decorative_assets[: max_assets], start=1):
            box = placements[index - 1]
            box, _adjusted = self._adapt_box_away_from_reserved(
                box,
                reserved_boxes=list(reserved_boxes or []),
                canvas_width=width,
                canvas_height=height,
                padding=18,
            )
            if not self._paste_decorative_asset(canvas, asset, box):
                continue
            zones.append(self._zone_manifest(f"decorative_{index}", asset.asset_role, box))
        return zones

    @staticmethod
    def _scene_text_value(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, list):
            return " ".join(str(item).strip() for item in value if str(item).strip())
        return str(value).strip()

    @classmethod
    def _scene_text_items(cls, value: object) -> list[str]:
        if value is None:
            return []
        if isinstance(value, list):
            raw_items = [str(item).strip() for item in value if str(item).strip()]
        else:
            text = str(value).strip()
            raw_items = re.split(
                r"(?:\r?\n|[;|])+|\s*•\s*|\s*[-*]\s+|(?:✔️|✓|✅|➜|➡️|âœ”ï¸|â€¢)\s*",
                text,
            )
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in raw_items:
            text = cls.SCENE_TEXT_SYMBOL_PATTERN.sub(" ", item)
            text = " ".join(text.split()).strip(" ,.;:-")
            if not text:
                continue
            key = text.casefold()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(text)
        return cleaned

    @staticmethod
    def _scene_graph_box(element: SceneGraphElement, width: int, height: int) -> tuple[int, int, int, int] | None:
        geometry = element.geometry
        if geometry.width is None or geometry.height is None:
            return None

        def _resolve(value: float | int | None, scale: int) -> int:
            if value is None:
                return 0
            numeric = float(value)
            if geometry.units == "normalized" and 0 <= numeric <= 1:
                return max(int(round(numeric * scale)), 0)
            return max(int(round(numeric)), 0)

        resolved_width = _resolve(geometry.width, width)
        resolved_height = _resolve(geometry.height, height)
        resolved_x = _resolve(geometry.x, width) if geometry.x is not None else 0
        resolved_y = _resolve(geometry.y, height) if geometry.y is not None else 0
        anchor_defaults = {
            "top-left": (0.08, 0.08),
            "top-center": (0.5, 0.08),
            "top-right": (0.82, 0.08),
            "center-left": (0.08, 0.46),
            "center": (0.5, 0.46),
            "center-right": (0.82, 0.46),
            "bottom-left": (0.08, 0.84),
            "bottom-center": (0.5, 0.84),
            "bottom-right": (0.82, 0.84),
        }
        if geometry.anchor and (geometry.x is None or geometry.y is None):
            anchor_x, anchor_y = anchor_defaults.get(geometry.anchor, (0.08, 0.08))
            if geometry.x is None:
                resolved_x = max(int(round((anchor_x * width) - (resolved_width / 2))), 0)
            if geometry.y is None:
                resolved_y = max(int(round((anchor_y * height) - (resolved_height / 2))), 0)
        return (
            resolved_x,
            resolved_y,
            min(resolved_x + resolved_width, width),
            min(resolved_y + resolved_height, height),
        )

    def _scene_graph_color(
        self,
        payload: RendererInput,
        style: dict,
        key: str,
        default: tuple[int, int, int],
        palette: dict[str, str],
    ) -> tuple[int, int, int]:
        explicit = str(style.get(key) or "").strip()
        if explicit:
            return self._parse_color(explicit, default)
        role_value = str(style.get(f"{key}_role") or style.get("fill_role") or "").strip().lower()
        if role_value:
            if role_value == "light_text":
                return (255, 255, 255)
            if role_value == "secondary_text":
                explicit_secondary = palette.get("secondary_text")
                if explicit_secondary:
                    return self._parse_color(explicit_secondary, default)
                return default
            if role_value == "primary_text":
                explicit_primary = palette.get("primary") or palette.get("headline")
                if explicit_primary:
                    return self._parse_color(explicit_primary, default)
                return default
            if role_value == "surface":
                explicit_surface = palette.get("surface") or palette.get("background")
                if explicit_surface:
                    return self._parse_color(explicit_surface, default)
                return default
            if role_value == "background":
                explicit_background = palette.get("background") or palette.get("surface")
                if explicit_background:
                    return self._parse_color(explicit_background, default)
                return default
            return self._parse_color(palette.get(role_value), default)
        return default

    def _resolve_scene_asset_path(self, payload: RendererInput, element: SceneGraphElement) -> str | None:
        def path_exists(storage_path: str) -> bool:
            if not str(storage_path or "").strip():
                return False
            if hasattr(self.storage, "exists"):
                return bool(self.storage.exists(storage_path))
            try:
                return Path(self.storage.absolute_path(storage_path)).exists()
            except Exception:
                return False

        if element.asset and element.asset.storage_path:
            candidate_path = str(element.asset.storage_path).strip()
            if candidate_path and path_exists(candidate_path):
                return candidate_path
        if element.role == "logo" and payload.logo_asset_path and path_exists(payload.logo_asset_path):
            return payload.logo_asset_path
        asset_role = (element.asset.asset_role if element.asset else "") or ""
        asset_id = str(element.asset.asset_id) if element.asset and element.asset.asset_id else ""
        candidates = [*payload.image_assets, *payload.decorative_assets]
        for asset in candidates:
            if not path_exists(asset.storage_path):
                continue
            if asset_id and str(asset.asset_id) == asset_id:
                return asset.storage_path
            if asset_role and str(asset.asset_role) == asset_role:
                return asset.storage_path
        for asset in payload.scene_graph.assets if payload.scene_graph else []:
            storage_path = str(asset.storage_path or "").strip()
            if not storage_path or not path_exists(storage_path):
                continue
            if asset_id and str(asset.asset_id or "") == asset_id:
                return storage_path
            if asset_role and str(asset.asset_role or "") == asset_role:
                return storage_path
        if asset_role in {"template_background", "template_preview"}:
            return payload.template_asset_path
        return None

    def _render_scene_graph(
        self,
        payload: RendererInput,
        size: dict[str, int],
    ) -> tuple[Image.Image, dict[str, object]]:
        scene_graph = payload.scene_graph
        assert scene_graph is not None
        palette = self._resolve_palette_roles(payload)
        primary = self._resolve_primary_color(payload, palette)
        accent = self._resolve_accent_color(payload, palette, primary)
        background = self._resolve_background_color(payload, palette, primary, accent)
        primary = self._ensure_text_contrast(primary, background)
        secondary_text = self._ensure_text_contrast(
            self._blend_color(primary, (255, 255, 255), 0.24),
            background,
            fallback=primary,
        )
        light_text = (255, 255, 255)

        bg_element = next((element for element in scene_graph.elements if element.visible and element.role == "background"), None)
        gradient_spec = None
        if bg_element:
            style = bg_element.style or {}
            if style.get("gradient_from") or style.get("gradient_to"):
                gradient_spec = {
                    "type": "linear",
                    "direction": str(style.get("gradient_direction") or "vertical"),
                    "start_color": style.get("gradient_from") or style.get("primary_fill") or palette.get("background"),
                    "end_color": style.get("gradient_to") or palette.get("primary"),
                }
            background = self._scene_graph_color(payload, style, "primary_fill", background, palette)
        using_base_canvas = bool(payload.base_canvas_asset_path)
        image = self._build_background_canvas(
            size["width"],
            size["height"],
            background,
            gradient_spec,
            base_canvas_asset_path=payload.base_canvas_asset_path,
        )
        template_rendered = False
        creative_mode = str((payload.creative_decision or {}).get("layout_mode") or scene_graph.layout_mode or "")
        template_policy = str(scene_graph.validation_hints.get("template_surface_policy") or "")
        if (
            not using_base_canvas
            and payload.template_asset_path
            and creative_mode in {"exact_template", "adapted_template"}
            and template_policy != "style_reference_only"
        ):
            template_rendered = self._paste_template_background(image, payload.template_asset_path)
        draw = ImageDraw.Draw(image)
        zones_used: list[dict[str, int | str | None]] = []
        text_blocks_used: list[dict[str, object]] = []
        logo_rendered = False
        image_rendered = False
        decorative_rendered = False
        image_assessments: list[dict[str, object]] = []
        text_fit_manifest: list[dict[str, object]] = []
        asset_boxes: list[dict[str, object]] = []
        overlap_checks: list[dict[str, object]] = []
        occupied_text_boxes: list[tuple[int, int, int, int]] = []
        reserved_text_roles = {"headline", "supporting_line", "body", "proof_points", "cta", "footer", "legal"}
        reserved_text_boxes = [
            box
            for element in scene_graph.elements
            if element.visible and element.role in reserved_text_roles
            for box in [self._scene_graph_box(element, size["width"], size["height"])]
            if box
        ]

        scene_layers = list(scene_graph.layers or [])
        scene_layer_set = set(scene_layers)

        def effective_layer_for(element: SceneGraphElement) -> str:
            explicit_layer = str(element.layer or "").strip()
            if explicit_layer:
                return explicit_layer
            role_name = str(element.role or "").strip().lower()
            element_type = str(element.element_type or "").strip().lower()
            preferred_by_role = {
                "background": ["background"],
                "image": ["hero_visuals", "primary_visual", "content"],
                "hero_visual": ["hero_visuals", "primary_visual", "content"],
                "icon": ["hero_visuals", "primary_visual", "content"],
                "headline": ["headline_overlay", "content"],
                "header": ["headline_overlay", "content"],
                "supporting_line": ["body_text_overlay", "content"],
                "body": ["body_text_overlay", "content"],
                "proof_points": ["body_text_overlay", "content"],
                "proof_point": ["body_text_overlay", "content"],
                "stat_highlights": ["body_text_overlay", "content"],
                "stat_highlight": ["body_text_overlay", "content"],
                "cta": ["cta_overlay", "footer", "content"],
                "logo": ["logo_overlay", "brand", "content"],
                "legal": ["footer", "cta_overlay", "content"],
                "footer": ["footer", "cta_overlay", "content"],
                "disclaimer": ["footer", "cta_overlay", "content"],
            }
            if element_type in {"image", "icon"} and role_name not in preferred_by_role:
                role_name = element_type
            if element_type == "logo":
                role_name = "logo"
            for candidate in preferred_by_role.get(role_name, []):
                if candidate in scene_layer_set:
                    return candidate
            return "content"

        elements_by_layer: dict[str, list[SceneGraphElement]] = {layer: [] for layer in scene_layers}
        for element in scene_graph.elements:
            if not element.visible:
                continue
            elements_by_layer.setdefault(effective_layer_for(element), []).append(element)

        ordered_layers = scene_layers + [layer for layer in elements_by_layer if layer not in scene_layer_set]
        for layer in ordered_layers:
            for element in elements_by_layer.get(layer, []):
                box = self._scene_graph_box(element, size["width"], size["height"])
                if not box:
                    continue
                style = element.style or {}
                role = element.role
                if element.element_type in {"background_shape", "decorative_shape"}:
                    if using_base_canvas:
                        continue
                    adapted_box, adjusted = self._adapt_box_away_from_reserved(
                        box,
                        reserved_boxes=occupied_text_boxes or reserved_text_boxes,
                        canvas_width=size["width"],
                        canvas_height=size["height"],
                        padding=18,
                    )
                    fill = self._scene_graph_color(payload, style, "fill", self._blend_color(background, accent, 0.18), palette)
                    shape = str(style.get("shape") or "rounded_rect")
                    if shape == "ellipse":
                        draw.ellipse(adapted_box, fill=fill)
                    else:
                        self._draw_panel(draw, adapted_box, fill, radius=int(style.get("border_radius") or 28))
                    decorative_rendered = True
                    asset_boxes.append({"role": role, "box": adapted_box, "adjusted": adjusted})
                    overlap_checks.append({"role": role, "passed": not self._box_overlaps_any(adapted_box, occupied_text_boxes or reserved_text_boxes, padding=18)})
                elif role == "logo" or element.element_type == "logo":
                    adapted_box, adjusted = self._adapt_box_away_from_reserved(
                        box,
                        reserved_boxes=occupied_text_boxes,
                        canvas_width=size["width"],
                        canvas_height=size["height"],
                        padding=20,
                    )
                    zone = type("Zone", (), self._zone_manifest(element.element_id, role, adapted_box))()
                    logo_path = self._resolve_scene_asset_path(payload, element)
                    logo_rendered = self._render_logo_box(
                        canvas=image,
                        draw=draw,
                        payload=payload,
                        box=adapted_box,
                        primary=primary,
                        accent=accent,
                        fill=primary,
                        logo_path=logo_path,
                    )
                    zones_used.append(self._zone_manifest(element.element_id, role, adapted_box))
                    asset_boxes.append({"role": role, "box": adapted_box, "adjusted": adjusted})
                    overlap_checks.append({"role": role, "passed": not self._box_overlaps_any(adapted_box, occupied_text_boxes, padding=20)})
                elif role in {"image", "icon"} or element.element_type in {"image", "icon"}:
                    adapted_box, adjusted = self._adapt_box_away_from_reserved(
                        box,
                        reserved_boxes=reserved_text_boxes,
                        canvas_width=size["width"],
                        canvas_height=size["height"],
                        padding=18,
                    )
                    if using_base_canvas:
                        zones_used.append(self._zone_manifest(element.element_id, role, adapted_box))
                        image_rendered = True
                        asset_boxes.append({"role": role, "box": adapted_box, "adjusted": adjusted})
                        overlap_checks.append({"role": role, "passed": not self._box_overlaps_any(adapted_box, reserved_text_boxes, padding=18)})
                        continue
                    asset_path = self._resolve_scene_asset_path(payload, element)
                    zone = type("Zone", (), self._zone_manifest(element.element_id, role, adapted_box))()
                    if asset_path:
                        visual_result = self._paste_visual(
                            image,
                            asset_path,
                            zone,
                            frame_fill=self._blend_color(background, (255, 255, 255), 0.72),
                            radius=int(style.get("border_radius") or 24),
                        )
                        image_rendered = bool(visual_result.get("rendered")) or image_rendered
                        if visual_result.get("rendered"):
                            image_assessments.append(visual_result)
                    else:
                        self._draw_panel(draw, adapted_box, self._blend_color(background, primary, 0.08), radius=int(style.get("border_radius") or 24))
                    zones_used.append(self._zone_manifest(element.element_id, role, adapted_box))
                    asset_boxes.append({"role": role, "box": adapted_box, "adjusted": adjusted})
                    overlap_checks.append({"role": role, "passed": not self._box_overlaps_any(adapted_box, reserved_text_boxes, padding=18)})
                elif role == "proof_points":
                    items = self._scene_text_items(element.text)
                    if not items:
                        continue
                    # Extract badge_style from validation_hints
                    validation_hints = element.validation_hints or {}
                    badge_style = validation_hints.get("badge_style")
                    proof_zones = self._draw_bullet_list(
                        draw,
                        items,
                        box,
                        secondary_text,
                        accent,
                        base_size=int(style.get("font_size") or 20),
                        family_hint=str(style.get("font_family") or "").strip() or None,
                        badge_style=badge_style,
                    )
                    zones_used.extend(proof_zones)
                    text_blocks_used.append({"role": role, "text": items})
                    occupied_text_boxes.extend(
                        (
                            int(zone["x"]),
                            int(zone["y"]),
                            int(zone["x"]) + int(zone["width"]),
                            int(zone["y"]) + int(zone["height"]),
                        )
                        for zone in proof_zones
                    )
                else:
                    text_value = self._scene_text_value(element.text)
                    if not text_value:
                        continue
                    fill = self._scene_graph_color(payload, style, "fill", primary if role == "headline" else secondary_text, palette)
                    background_fill = None
                    if style.get("background_fill") or style.get("background_fill_role"):
                        background_fill = self._scene_graph_color(payload, style, "background_fill", primary, palette)
                    overlay_minimal = using_base_canvas and role not in {"cta"}
                    zone_manifest = self._zone_manifest(
                        element.element_id,
                        role,
                        box,
                        int(style.get("max_lines")) if style.get("max_lines") else None,
                    )
                    zone = type("Zone", (), zone_manifest)()
                    padding = int(style.get("padding") or (18 if role == "cta" else 0))
                    if overlay_minimal:
                        cleanup_fill, fill, cleanup_padding = self._base_canvas_text_cleanup_fill(
                            fill=fill,
                            background=background,
                            primary=primary,
                            role=role,
                        )
                        background_fill = background_fill or cleanup_fill
                        padding = max(padding, cleanup_padding)
                    fit_manifest = self._draw_text_block(
                        draw,
                        text_value,
                        zone,
                        fill,
                        int(style.get("font_size") or (56 if role == "headline" else 22)),
                        align=str(style.get("align") or "left"),
                        background_fill=background_fill,
                        padding=padding,
                        style_hint=style,
                        font_family=str(style.get("font_family") or "").strip() or None,
                    )
                    zones_used.append(zone_manifest)
                    text_blocks_used.append({"role": role, "text": text_value})
                    if fit_manifest.get("occupied_box"):
                        occupied_text_boxes.append(tuple(fit_manifest["occupied_box"]))
                    text_fit_manifest.append(fit_manifest)

        return image, {
            "logo_rendered": logo_rendered,
            "image_rendered": image_rendered,
            "template_rendered": template_rendered,
            "decorative_rendered": decorative_rendered,
            "render_path": "scene_graph_direct",
            "layout_variant": f"scene_graph_{scene_graph.styles.get('layout_archetype') or scene_graph.styles.get('layout_type') or scene_graph.layout_mode}",
            "zones_used": zones_used,
            "text_blocks_used": text_blocks_used,
            "image_assessments": image_assessments,
            "text_fit": text_fit_manifest,
            "asset_boxes": asset_boxes,
            "overlap_checks": overlap_checks,
            "pre_shortening": {
                "headline": {"original_length": len(str(payload.text.headline or "")), "fitted_length": len(str(payload.text.headline or "")), "shortened": False},
                "body": {"original_length": len(str(payload.text.body or "")), "fitted_length": len(str(payload.text.body or "")), "shortened": False},
            },
            "content_structure_type": (
                ((payload.text.metadata or {}).get("content_structure", {}) if isinstance(payload.text.metadata, dict) else {}).get("creative_type")
                or payload.blueprint.layout_type
            ),
        }

    def _save_image_asset(self, tenant_id: UUID, brand_space_id: UUID, filename: str, image: Image.Image, role: str) -> dict:
        buffer = BytesIO()
        format_name = "PNG" if filename.lower().endswith(".png") else "JPEG"
        image.save(buffer, format=format_name)
        stored = self.storage.save_bytes(tenant_id, brand_space_id, "generated", filename, buffer.getvalue())
        return {
            "asset_id": uuid4(),
            "mime_type": "image/png" if format_name == "PNG" else "image/jpeg",
            "storage_path": stored.storage_path,
            "width": image.width,
            "height": image.height,
            "asset_role": role,
        }

    @staticmethod
    def _image_to_pdf_bytes(image: Image.Image) -> bytes:
        buffer = BytesIO()
        image.save(buffer, format="PDF")
        return buffer.getvalue()

    def _document_export(self, payload: RendererInput, image: Image.Image):
        document = Document()
        document.add_heading(payload.text.headline, level=1)
        document.add_paragraph(payload.text.body)
        if payload.text.cta:
            document.add_paragraph(payload.text.cta)
        temp = BytesIO()
        image.save(temp, format="PNG")
        temp.seek(0)
        document.add_picture(temp)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return self.storage.save_bytes(
            payload.tenant_id,
            payload.brand_space_id,
            "generated",
            f"export-{payload.content_version_id}.docx",
            buffer.getvalue(),
        )

    def _document_export_multi(self, payload: RendererInput, pages: list[Image.Image]):
        document = Document()
        document.add_heading(payload.text.headline, level=1)
        document.add_paragraph(payload.text.body)
        if payload.text.cta:
            document.add_paragraph(payload.text.cta)
        for page in pages:
            temp = BytesIO()
            page.save(temp, format="PNG")
            temp.seek(0)
            document.add_picture(temp)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return self.storage.save_bytes(
            payload.tenant_id,
            payload.brand_space_id,
            "generated",
            f"export-{payload.content_version_id}.docx",
            buffer.getvalue(),
        )

    def _render_page(
        self,
        payload: RendererInput,
        size: dict[str, int],
        page_text: dict,
    ) -> tuple[Image.Image, dict[str, object]]:
        palette = self._resolve_palette_roles(payload)
        primary = self._resolve_primary_color(payload, palette)
        accent = self._resolve_accent_color(payload, palette, primary)
        background = self._resolve_background_color(payload, palette, primary, accent)
        primary = self._ensure_text_contrast(primary, background)
        secondary_text = self._ensure_text_contrast(
            self._parse_color(
            payload.brand_visual_rules.get("body_color") or palette.get("secondary_text"),
            self._blend_color(background, primary, 0.78),
            ),
            background,
            fallback=primary,
        )
        light_text = (255, 255, 255)

        gradient_spec = self._resolve_gradient_spec(payload, background, primary, accent)
        image = self._build_background_canvas(size["width"], size["height"], background, gradient_spec)
        draw = ImageDraw.Draw(image)
        style_reference_only = self._template_is_style_reference_only(payload)
        template_rendered = False
        if not style_reference_only:
            template_rendered = self._paste_template_background(image, payload.template_asset_path)
        if template_rendered:
            draw = ImageDraw.Draw(image)

        preset = payload.studio_panel.get("platform_preset", payload.blueprint.platform_preset)
        prefer_blueprint_zones = self._prefer_blueprint_zone_rendering(payload)
        if payload.blueprint.layout_type == "infographic" and not template_rendered and not prefer_blueprint_zones:
            return self._render_infographic_page(
                payload=payload,
                size=size,
                page_text=page_text,
                preset=preset,
                background=background,
                primary=primary,
                accent=accent,
                secondary_text=secondary_text,
                light_text=light_text,
            )
        if (
            payload.blueprint.layout_type in {"static", "carousel", "pdf"}
            and not template_rendered
            and preset in {"instagram", "linkedin", "x", "youtube_thumbnail"}
            and not prefer_blueprint_zones
        ):
            return self._render_social_page(
                payload=payload,
                size=size,
                page_text=page_text,
                preset=preset,
                background=background,
                primary=primary,
                accent=accent,
                secondary_text=secondary_text,
                light_text=light_text,
            )

        zone_map = {zone.role: zone for zone in payload.blueprint.zones}
        image_rendered = False
        image_assessments: list[dict[str, object]] = []
        text_fit_manifest: list[dict[str, object]] = []
        asset_boxes: list[dict[str, object]] = []
        overlap_checks: list[dict[str, object]] = []
        reserved_text_boxes = [
            self._zone_box(zone)
            for role, zone in zone_map.items()
            if role in {"header", "headline", "body", "proof_points", "proof_point", "stat_highlights", "stat_highlight", "cta", "footer"}
        ]
        adaptation_plan = (payload.layout_decision or {}).get("adaptation_plan", {}) or {}
        allow_generated_image_overlay = (
            not template_rendered
            or bool(adaptation_plan.get("visual_slot_synthesis"))
            or bool(adaptation_plan.get("replace_primary_visual"))
        )
        if allow_generated_image_overlay and payload.image_assets and page_text.get("show_image"):
            for index, image_zone_entry in enumerate(payload.blueprint.image_zones):
                zone_id = image_zone_entry["zone_id"]
                zone_role = str(image_zone_entry.get("role") or "").strip().lower()
                image_zone = next((item for item in payload.blueprint.zones if item.zone_id == zone_id), None)
                if not image_zone and zone_role:
                    image_zone = next(
                        (
                            item
                            for item in payload.blueprint.zones
                            if str(item.role or "").strip().lower() in {zone_role, "image"}
                        ),
                        None,
                    )
                if not image_zone:
                    continue
                adapted_box, adjusted = self._adapt_box_away_from_reserved(
                    self._zone_box(image_zone),
                    reserved_boxes=reserved_text_boxes,
                    canvas_width=size["width"],
                    canvas_height=size["height"],
                    padding=18,
                )
                image_zone = type("Zone", (), self._zone_manifest(image_zone.zone_id, image_zone.role, adapted_box, image_zone.max_lines))()
                asset = payload.image_assets[min(index, len(payload.image_assets) - 1)]
                visual_result = self._paste_visual(
                    image,
                    asset.storage_path,
                    image_zone,
                    frame_fill=self._blend_color(background, (255, 255, 255), 0.68),
                    radius=22,
                )
                image_rendered = bool(visual_result.get("rendered")) or image_rendered
                if visual_result.get("rendered"):
                    image_assessments.append(visual_result)
                asset_boxes.append({"role": "image", "box": adapted_box, "adjusted": adjusted})
                overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(adapted_box, reserved_text_boxes, padding=18)})

        logo_zone = zone_map.get("logo")
        logo_rendered = False
        should_overlay_logo = bool(logo_zone) and (
            not template_rendered
            or bool(adaptation_plan.get("logo_injection_required"))
            or bool(adaptation_plan.get("replace_logo"))
        )
        if should_overlay_logo and logo_zone:
            adapted_logo_box, adjusted = self._adapt_box_away_from_reserved(
                self._zone_box(logo_zone),
                reserved_boxes=reserved_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=20,
            )
            logo_zone = type("Zone", (), self._zone_manifest(logo_zone.zone_id, logo_zone.role, adapted_logo_box, logo_zone.max_lines))()
            logo_rendered = self._paste_logo(image, payload.logo_asset_path, logo_zone)
            asset_boxes.append({"role": "logo", "box": adapted_logo_box, "adjusted": adjusted})
            overlap_checks.append({"role": "logo", "passed": not self._box_overlaps_any(adapted_logo_box, reserved_text_boxes, padding=20)})

        header_zone = zone_map.get("header")
        headline_zone = zone_map.get("headline")
        body_zone = zone_map.get("body")
        proof_zone = zone_map.get("proof_points") or zone_map.get("proof_point")
        stat_zone = zone_map.get("stat_highlights") or zone_map.get("stat_highlight")
        cta_zone = zone_map.get("cta")
        footer_zone = zone_map.get("footer")
        composition = self._composition_plan(payload)
        text_style_plan = composition.get("text_style_plan", {}) if isinstance(composition, dict) else {}

        if header_zone:
            text_fit_manifest.append(self._draw_text_block(
                draw=draw,
                text=page_text.get("header") or payload.text.metadata.get("header", ""),
                zone=header_zone,
                fill=primary,
                base_size=26,
                padding=12,
                style_hint=self._zone_style_hint(payload, "header"),
            ))
        if headline_zone:
            text_fit_manifest.append(self._draw_text_block(
                draw=draw,
                text=page_text.get("headline", ""),
                zone=headline_zone,
                fill=primary,
                base_size=52,
                padding=16,
                style_hint=self._zone_style_hint(payload, "headline"),
            ))
        if body_zone:
            text_fit_manifest.append(self._draw_text_block(
                draw=draw,
                text=page_text.get("body_display") or page_text.get("body", ""),
                zone=body_zone,
                fill=secondary_text,
                base_size=28,
                padding=14,
                style_hint=self._zone_style_hint(payload, "body"),
            ))
        proof_points = self._page_list_value(page_text.get("proof_points"))
        if proof_zone and proof_points:
            # Extract badge_style from scene graph if available
            badge_style = None
            if payload.scene_graph and hasattr(payload.scene_graph, 'elements'):
                proof_element = next(
                    (el for el in payload.scene_graph.elements if el.role == "proof_points"),
                    None
                )
                if proof_element and proof_element.validation_hints:
                    badge_style = proof_element.validation_hints.get("badge_style")

            proof_zones = self._draw_bullet_list(
                draw,
                proof_points,
                self._zone_box(proof_zone),
                secondary_text,
                accent,
                base_size=20,
                badge_style=badge_style,
            )
            occupied_text_boxes = [
                tuple(item["occupied_box"])
                for item in text_fit_manifest
                if isinstance(item, dict) and item.get("occupied_box")
            ]
            overlap_checks.append(
                {
                    "role": "proof_points",
                    "passed": all(
                        not self._box_overlaps_any(
                            (
                                int(zone["x"]),
                                int(zone["y"]),
                                int(zone["x"]) + int(zone["width"]),
                                int(zone["y"]) + int(zone["height"]),
                            ),
                            occupied_text_boxes,
                            padding=12,
                        )
                        for zone in proof_zones
                    ),
                }
            )
            asset_boxes.extend(
                {"role": "proof_point", "box": (
                    int(zone["x"]),
                    int(zone["y"]),
                    int(zone["x"]) + int(zone["width"]),
                    int(zone["y"]) + int(zone["height"]),
                ), "adjusted": False}
                for zone in proof_zones
            )
        stat_highlights = self._page_list_value(page_text.get("stat_highlights"))
        if stat_zone and stat_highlights:
            stat_zones = self._draw_stat_cards(
                draw,
                stat_highlights,
                self._zone_box(stat_zone),
                primary,
                accent,
                background,
            )
            asset_boxes.extend(
                {"role": "stat_highlight", "box": (
                    int(zone["x"]),
                    int(zone["y"]),
                    int(zone["x"]) + int(zone["width"]),
                    int(zone["y"]) + int(zone["height"]),
                ), "adjusted": False}
                for zone in stat_zones
            )
        if cta_zone and page_text.get("cta"):
            cta_style_hint = self._zone_style_hint(payload, "cta")
            cta_fill = light_text
            cta_background_fill = primary
            template_native_cta = template_rendered and (
                str(text_style_plan.get("cta_style") or "") == "template_native"
                or payload.blueprint.source_mode in {"exact_template", "adapted_template"}
            )
            if template_native_cta and not (isinstance(cta_style_hint, dict) and cta_style_hint.get("background_color")):
                cta_fill = self._resolved_fill_from_style_hint(cta_style_hint, accent)
                cta_background_fill = primary
            text_fit_manifest.append(self._draw_text_block(
                draw=draw,
                text=page_text.get("cta", ""),
                zone=cta_zone,
                fill=cta_fill,
                base_size=24,
                align=payload.blueprint.cta_placement.get("alignment", "left"),
                background_fill=cta_background_fill,
                padding=18,
                style_hint=cta_style_hint,
            ))
        if footer_zone:
            footer_text = self._footer_text(payload, page_text)
            if footer_text:
                text_fit_manifest.append(self._draw_text_block(
                    draw=draw,
                    text=footer_text,
                    zone=footer_zone,
                    fill=secondary_text,
                    base_size=20,
                    padding=12,
                    style_hint=self._zone_style_hint(payload, "footer"),
                ))
        return image, {
            "logo_rendered": logo_rendered,
            "image_rendered": image_rendered,
            "template_rendered": template_rendered,
            "image_assessments": image_assessments,
            "render_path": "page_generic",
            "layout_variant": "generic_zones",
            "zones_used": [zone.model_dump() for zone in payload.blueprint.zones],
            "text_blocks_used": [
                {"role": "headline", "text": page_text.get("headline", "")},
                {"role": "body", "text": page_text.get("body_display") or page_text.get("body", "")},
                {"role": "proof_points", "text": proof_points},
                {"role": "stat_highlights", "text": stat_highlights},
                {"role": "cta", "text": page_text.get("cta", "")},
                {"role": "footer", "text": self._footer_text(payload, page_text)},
            ],
            "text_fit": text_fit_manifest,
            "asset_boxes": asset_boxes,
            "overlap_checks": overlap_checks,
            "content_structure_type": page_text.get("content_structure_type") or payload.blueprint.layout_type,
            "pre_shortening": {
                "headline": {
                    "original_length": len(str(page_text.get("headline", "") or "")),
                    "fitted_length": len(str(page_text.get("headline", "") or "")),
                    "shortened": False,
                },
                "body": {
                    "original_length": len(str(page_text.get("body", "") or "")),
                    "fitted_length": len(str(page_text.get("body_display") or page_text.get("body", "") or "")),
                    "shortened": str(page_text.get("body_display") or page_text.get("body", "") or "").strip() != str(page_text.get("body", "") or "").strip(),
                },
            },
        }

    def _render_social_page(
        self,
        payload: RendererInput,
        size: dict[str, int],
        page_text: dict,
        preset: str,
        background: tuple[int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        secondary_text: tuple[int, int, int],
        light_text: tuple[int, int, int],
    ) -> tuple[Image.Image, dict[str, object]]:
        composition = self._composition_plan(payload)
        layout_plan = composition.get("layout_plan", {}) if isinstance(composition, dict) else {}
        layout_archetype = str(payload.blueprint.layout_archetype or layout_plan.get("layout_archetype") or "")
        gradient_spec = self._resolve_gradient_spec(payload, background, primary, accent)
        image = self._build_background_canvas(size["width"], size["height"], background, gradient_spec)
        draw = ImageDraw.Draw(image)
        palette = self._resolve_palette_roles(payload)
        color_tokens = self._reference_palette_tokens(
            payload,
            palette=palette,
            background=background,
            primary=primary,
            accent=accent,
        )
        background = color_tokens["background"]
        primary = color_tokens["primary"]
        accent = color_tokens["accent"]
        accent_soft = color_tokens["accent_soft"]
        stroke = color_tokens["stroke"]
        surface = color_tokens["surface"]
        proof_surface = color_tokens["proof_surface"]
        image_frame_fill = color_tokens["image_frame_fill"]
        brand_name = self._brand_name(payload)
        badges = self._supporting_badges(payload, preset)
        section_label = self._section_label(payload, preset)

        headline_text = str(page_text.get("headline_display") or "").strip() or self._headline_copy(page_text.get("headline", ""), preset)
        body_text = str(page_text.get("body_display") or "").strip() or self._social_body_copy(page_text.get("body", ""), preset)
        supporting_line = str(page_text.get("supporting_line") or "").strip() or self._supporting_line(payload, preset, page_text.get("body", ""))
        proof_points = self._page_list_value(page_text.get("proof_points")) or self._proof_points(payload, preset, page_text.get("body", ""))
        stat_highlights = self._page_list_value(page_text.get("stat_highlights")) or self._stat_highlights(payload, preset)
        cta_text = str(page_text.get("cta_display") or page_text.get("cta", "") or "").strip()
        image_assessments: list[dict[str, object]] = []
        text_fit_manifest: list[dict[str, object]] = []
        asset_boxes: list[dict[str, object]] = []
        overlap_checks: list[dict[str, object]] = []
        pre_shortening = {
            "headline": {
                "original_length": len(str(page_text.get("headline", "") or "")),
                "fitted_length": len(headline_text),
                "shortened": str(page_text.get("headline", "") or "").strip() != headline_text.strip(),
            },
            "body": {
                "original_length": len(str(page_text.get("body", "") or "")),
                "fitted_length": len(body_text),
                "shortened": str(page_text.get("body", "") or "").strip() != body_text.strip(),
            },
        }

        def _occupied_from_fit(items: list[dict[str, object]]) -> list[tuple[int, int, int, int]]:
            return [
                tuple(item["occupied_box"])
                for item in items
                if isinstance(item, dict) and item.get("occupied_box")
            ]

        if preset == "instagram" and layout_archetype == "checklist_card":
            margin = max(52, int(size["width"] * 0.06))
            outer_box = (margin, margin, size["width"] - margin, size["height"] - margin)
            header_box = (outer_box[0] + 28, outer_box[1] + 26, outer_box[2] - 28, outer_box[1] + 116)
            headline_box = (outer_box[0] + 28, outer_box[1] + 120, outer_box[2] - 28, outer_box[1] + 250)
            support_box = (outer_box[0] + 28, outer_box[1] + 258, outer_box[2] - 28, outer_box[1] + 322)
            proof_box = (outer_box[0] + 28, outer_box[1] + 340, outer_box[2] - 28, outer_box[1] + 592)
            image_box = (outer_box[0] + 28, outer_box[1] + 610, outer_box[2] - 28, outer_box[1] + 836)
            cta_box = (outer_box[0] + 28, outer_box[3] - 112, outer_box[0] + 360, outer_box[3] - 26)

            self._draw_panel(draw, outer_box, surface, radius=36, outline=stroke, width=2)
            self._draw_panel(draw, (outer_box[0] + 14, outer_box[1] - 8, outer_box[2] - 110, outer_box[1] + 86), accent_soft, radius=26)
            self._draw_section_label(draw, section_label, header_box, primary, accent)

            headline_zone = self._zone_manifest("headline", "headline", headline_box, 3)
            support_zone = self._zone_manifest("supporting_line", "supporting_line", support_box, 2)
            headline_fit = self._draw_text_block(
                draw,
                headline_text,
                type("Zone", (), headline_zone)(),
                primary,
                54,
                padding=0,
                style_hint=self._zone_style_hint(payload, "headline"),
            )
            support_fit = self._draw_text_block(
                draw,
                supporting_line or body_text,
                type("Zone", (), support_zone)(),
                secondary_text,
                24,
                padding=0,
                style_hint=self._zone_style_hint(payload, "body"),
            )
            text_fit_manifest.extend([headline_fit, support_fit])
            occupied_text_boxes = _occupied_from_fit([headline_fit, support_fit])

            proof_panel = (proof_box[0] - 6, proof_box[1] - 8, proof_box[2] + 6, proof_box[3] + 10)
            self._draw_panel(draw, proof_panel, proof_surface, radius=26, outline=stroke, width=1)

            # Extract badge_style from scene graph if available
            badge_style = None
            if payload.scene_graph and hasattr(payload.scene_graph, 'elements'):
                proof_element = next(
                    (el for el in payload.scene_graph.elements if el.role == "proof_points"),
                    None
                )
                if proof_element and proof_element.validation_hints:
                    badge_style = proof_element.validation_hints.get("badge_style")

            proof_zones = self._draw_bullet_list(draw, proof_points[:4], proof_box, secondary_text, accent, base_size=22, badge_style=badge_style)
            occupied_text_boxes.extend(
                (
                    int(zone["x"]),
                    int(zone["y"]),
                    int(zone["x"]) + int(zone["width"]),
                    int(zone["y"]) + int(zone["height"]),
                )
                for zone in proof_zones
            )

            image_rendered = False
            image_box, image_adjusted = self._adapt_box_away_from_reserved(
                image_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=18,
            )
            if payload.image_assets and page_text.get("show_image"):
                visual_result = self._paste_visual_card(
                    image,
                    payload.image_assets[0].storage_path,
                    image_box,
                    radius=28,
                    frame_fill=image_frame_fill,
                    padding=0,
                )
                image_rendered = bool(visual_result.get("rendered"))
                if visual_result.get("rendered"):
                    image_assessments.append(visual_result)
            if not image_rendered:
                self._draw_editorial_motif(draw, image_box, primary, accent, background)
            asset_boxes.append({"role": "image", "box": image_box, "adjusted": image_adjusted})
            overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(image_box, occupied_text_boxes, padding=18)})

            logo_rendered = False
            logo_box = None
            logo_box = (outer_box[2] - 220, outer_box[1] + 28, outer_box[2] - 30, outer_box[1] + 96)
            logo_box, logo_adjusted = self._adapt_box_away_from_reserved(
                logo_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=20,
            )
            logo_rendered = self._render_logo_box(
                canvas=image,
                draw=draw,
                payload=payload,
                box=logo_box,
                primary=primary,
                accent=accent,
                fill=primary,
            )
            asset_boxes.append({"role": "logo", "box": logo_box, "adjusted": logo_adjusted})
            overlap_checks.append({"role": "logo", "passed": not self._box_overlaps_any(logo_box, occupied_text_boxes, padding=20)})

            if cta_text:
                text_fit_manifest.append(self._draw_text_block(
                    draw,
                    cta_text,
                    type("Zone", (), self._zone_manifest("cta", "cta", cta_box, 2))(),
                    light_text,
                    24,
                    background_fill=primary,
                    padding=18,
                ))

            zones_used = [
                self._zone_manifest("section_label", "section_label", header_box, 1),
                headline_zone,
                support_zone,
                *proof_zones,
                self._zone_manifest("image", "image", image_box),
                self._zone_manifest("logo", "logo", logo_box),
            ]
            if cta_text:
                zones_used.append(self._zone_manifest("cta", "cta", cta_box, 2))
            return image, {
                "logo_rendered": logo_rendered,
                "image_rendered": image_rendered,
                "template_rendered": False,
                "image_assessments": image_assessments,
                "render_path": "social_page",
                "layout_variant": "instagram_checklist_card",
                "zones_used": zones_used,
                "text_blocks_used": [
                    {"role": "brand", "text": brand_name},
                    {"role": "section_label", "text": section_label},
                    {"role": "headline", "text": headline_text},
                    {"role": "supporting_line", "text": supporting_line or body_text},
                    {"role": "proof_points", "text": proof_points[:4]},
                    {"role": "cta", "text": cta_text},
                ],
                "text_fit": text_fit_manifest,
                "asset_boxes": asset_boxes,
                "overlap_checks": overlap_checks,
                "pre_shortening": pre_shortening,
                "content_structure_type": page_text.get("content_structure_type") or "static",
            }

        if preset == "instagram":
            margin = max(52, int(size["width"] * 0.06))
            text_card = (margin, margin, size["width"] - margin, int(size["height"] * 0.4))
            image_box = (margin, int(size["height"] * 0.42), size["width"] - margin, size["height"] - 150)
            cta_box = (margin, image_box[3] + 18, margin + min(380, int(size["width"] * 0.42)), size["height"] - margin)

            self._draw_panel(draw, (margin + 14, margin - 10, size["width"] - margin - 120, 112), accent_soft, radius=28)
            self._draw_panel(draw, text_card, surface, radius=34, outline=stroke, width=2)
            self._draw_section_label(draw, section_label, (margin + 34, margin + 28, margin + 220, margin + 62), primary, accent)
            headline_zone = self._zone_manifest("headline", "headline", (text_card[0] + 34, text_card[1] + 78, text_card[2] - 34, text_card[1] + 184), 3)
            support_zone = self._zone_manifest("supporting_line", "supporting_line", (text_card[0] + 34, text_card[1] + 188, text_card[2] - 34, text_card[1] + 248), 2)
            cta_zone = self._zone_manifest("cta", "cta", cta_box, 2)
            headline_fit = self._draw_text_block(
                draw,
                headline_text,
                type("Zone", (), headline_zone)(),
                primary,
                58,
                padding=0,
                style_hint=self._zone_style_hint(payload, "headline"),
            )
            support_fit = self._draw_text_block(
                draw,
                supporting_line or body_text,
                type("Zone", (), support_zone)(),
                secondary_text,
                24,
                padding=0,
                style_hint=self._zone_style_hint(payload, "body"),
            )
            text_fit_manifest.extend([headline_fit, support_fit])
            occupied_text_boxes = _occupied_from_fit([headline_fit, support_fit])
            image_box, image_adjusted = self._adapt_box_away_from_reserved(
                image_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=18,
            )
            image_rendered = False
            if payload.image_assets and page_text.get("show_image"):
                visual_result = self._paste_visual_card(
                    image,
                    payload.image_assets[0].storage_path,
                    image_box,
                    radius=34,
                    frame_fill=image_frame_fill,
                    padding=0,
                )
                image_rendered = bool(visual_result.get("rendered"))
                if visual_result.get("rendered"):
                    image_assessments.append(visual_result)
            if not image_rendered:
                self._draw_editorial_motif(draw, image_box, primary, accent, background)
            asset_boxes.append({"role": "image", "box": image_box, "adjusted": image_adjusted})
            overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(image_box, occupied_text_boxes, padding=18)})
            if cta_text:
                text_fit_manifest.append(self._draw_text_block(
                    draw,
                    cta_text,
                    type("Zone", (), cta_zone)(),
                    light_text,
                    24,
                    background_fill=primary,
                    padding=18,
                ))
            logo_rendered = False
            logo_box = None
            logo_box = (text_card[2] - 180, text_card[1] + 24, text_card[2] - 28, text_card[1] + 92)
            logo_box, logo_adjusted = self._adapt_box_away_from_reserved(
                logo_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=20,
            )
            logo_rendered = self._render_logo_box(
                canvas=image,
                draw=draw,
                payload=payload,
                box=logo_box,
                primary=primary,
                accent=accent,
                fill=primary,
            )
            asset_boxes.append({"role": "logo", "box": logo_box, "adjusted": logo_adjusted})
            overlap_checks.append({"role": "logo", "passed": not self._box_overlaps_any(logo_box, occupied_text_boxes, padding=20)})
            badges_box = (text_card[0] + 34, text_card[3] - 70, text_card[2] - 34, text_card[3] - 22)
            proof_chips = badges + stat_highlights
            self._draw_badges(draw, proof_chips[:3], badges_box, primary, accent)
            zones_used = [
                self._zone_manifest("headline", "headline", (headline_zone["x"], headline_zone["y"], headline_zone["x"] + headline_zone["width"], headline_zone["y"] + headline_zone["height"]), 3),
                self._zone_manifest("supporting_line", "supporting_line", (support_zone["x"], support_zone["y"], support_zone["x"] + support_zone["width"], support_zone["y"] + support_zone["height"]), 2),
                self._zone_manifest("image", "image", image_box),
                cta_zone,
            ]
            zones_used.append(self._zone_manifest("logo", "logo", logo_box))
            zones_used.append(self._zone_manifest("section_label", "section_label", (margin + 34, margin + 28, margin + 220, margin + 62), 1))
            if proof_chips:
                zones_used.append(self._zone_manifest("badges", "badges", badges_box, 1))
            return image, {
                "logo_rendered": logo_rendered,
                "image_rendered": image_rendered,
                "template_rendered": False,
                "image_assessments": image_assessments,
                "render_path": "social_page",
                "layout_variant": "instagram_editorial",
                "zones_used": zones_used,
                "text_blocks_used": [
                    {"role": "brand", "text": brand_name},
                    {"role": "section_label", "text": section_label},
                    {"role": "headline", "text": headline_text},
                    {"role": "supporting_line", "text": supporting_line or body_text},
                    {"role": "cta", "text": cta_text},
                    {"role": "badges", "text": proof_chips[:3]},
                ],
                "text_fit": text_fit_manifest,
                "asset_boxes": asset_boxes,
                "overlap_checks": overlap_checks,
                "pre_shortening": pre_shortening,
                "content_structure_type": page_text.get("content_structure_type") or "static",
            }

        if preset == "linkedin":
            margin_x = max(48, int(size["width"] * 0.05))
            margin_y = max(30, int(size["height"] * 0.07))
            gap = max(24, int(size["width"] * 0.028))
            text_card = (margin_x, margin_y + 8, int(size["width"] * 0.48), size["height"] - margin_y)
            image_box = (text_card[2] + gap, margin_y + 48, size["width"] - margin_x, size["height"] - margin_y - 18)
            headline_box = (text_card[0] + 38, text_card[1] + 88, text_card[2] - 34, text_card[1] + 194)
            support_box = (text_card[0] + 38, text_card[1] + 198, text_card[2] - 34, text_card[1] + 252)
            proof_box = (text_card[0] + 38, text_card[1] + 270, text_card[2] - 38, text_card[1] + 408)
            stats_box = (text_card[0] + 38, text_card[1] + 420, text_card[2] - 38, text_card[1] + 494)
            cta_box = (text_card[0] + 38, text_card[3] - 98, text_card[0] + 336, text_card[3] - 28)

            self._draw_panel(draw, text_card, surface, radius=32, outline=stroke, width=2)
            self._draw_panel(draw, (image_box[0] - 12, image_box[1] - 16, image_box[2] + 12, image_box[3] + 12), accent_soft, radius=36)
            self._draw_section_label(draw, section_label, (text_card[0] + 38, text_card[1] + 28, text_card[0] + 220, text_card[1] + 62), primary, accent)
            headline_zone = self._zone_manifest("headline", "headline", headline_box, 3)
            support_zone = self._zone_manifest("supporting_line", "supporting_line", support_box, 2)
            cta_zone = self._zone_manifest("cta", "cta", cta_box, 2)
            headline_fit = self._draw_text_block(
                draw,
                headline_text,
                type("Zone", (), headline_zone)(),
                primary,
                46,
                padding=0,
                style_hint=self._zone_style_hint(payload, "headline"),
            )
            support_fit = self._draw_text_block(
                draw,
                supporting_line or body_text,
                type("Zone", (), support_zone)(),
                secondary_text,
                22,
                padding=0,
                style_hint=self._zone_style_hint(payload, "body"),
            )
            text_fit_manifest.extend([headline_fit, support_fit])
            occupied_text_boxes = _occupied_from_fit([headline_fit, support_fit])
            image_box, image_adjusted = self._adapt_box_away_from_reserved(
                image_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=18,
            )
            image_rendered = False
            if payload.image_assets and page_text.get("show_image"):
                visual_result = self._paste_visual_card(
                    image,
                    payload.image_assets[0].storage_path,
                    image_box,
                    radius=32,
                    frame_fill=surface,
                    padding=0,
                )
                image_rendered = bool(visual_result.get("rendered"))
                if visual_result.get("rendered"):
                    image_assessments.append(visual_result)
            if not image_rendered:
                self._draw_editorial_motif(draw, image_box, primary, accent, background)
            asset_boxes.append({"role": "image", "box": image_box, "adjusted": image_adjusted})
            overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(image_box, occupied_text_boxes, padding=18)})

            # Extract badge_style from scene graph if available
            badge_style = None
            if payload.scene_graph and hasattr(payload.scene_graph, 'elements'):
                proof_element = next(
                    (el for el in payload.scene_graph.elements if el.role == "proof_points"),
                    None
                )
                if proof_element and proof_element.validation_hints:
                    badge_style = proof_element.validation_hints.get("badge_style")

            proof_zones = self._draw_bullet_list(draw, proof_points, proof_box, secondary_text, accent, base_size=19, badge_style=badge_style)
            stat_zones = self._draw_stat_cards(draw, stat_highlights, stats_box, primary, accent, background)
            occupied_text_boxes.extend(
                (
                    int(zone["x"]),
                    int(zone["y"]),
                    int(zone["x"]) + int(zone["width"]),
                    int(zone["y"]) + int(zone["height"]),
                )
                for zone in [*proof_zones, *stat_zones]
            )
            if cta_text:
                text_fit_manifest.append(self._draw_text_block(
                    draw,
                    cta_text,
                    type("Zone", (), cta_zone)(),
                    light_text,
                    22,
                    background_fill=primary,
                    padding=16,
                ))
            logo_rendered = False
            logo_box = None
            logo_box = (image_box[2] - 178, image_box[1] - 30, image_box[2] - 24, image_box[1] + 28)
            logo_box, logo_adjusted = self._adapt_box_away_from_reserved(
                logo_box,
                reserved_boxes=occupied_text_boxes,
                canvas_width=size["width"],
                canvas_height=size["height"],
                padding=20,
            )
            logo_rendered = self._render_logo_box(
                canvas=image,
                draw=draw,
                payload=payload,
                box=logo_box,
                primary=primary,
                accent=accent,
                fill=primary,
            )
            asset_boxes.append({"role": "logo", "box": logo_box, "adjusted": logo_adjusted})
            overlap_checks.append({"role": "logo", "passed": not self._box_overlaps_any(logo_box, occupied_text_boxes, padding=20)})
            zones_used = [
                self._zone_manifest("section_label", "section_label", (text_card[0] + 38, text_card[1] + 28, text_card[0] + 220, text_card[1] + 62), 1),
                headline_zone,
                support_zone,
                self._zone_manifest("image", "image", image_box),
                *proof_zones,
                *stat_zones,
            ]
            if cta_text:
                zones_used.append(cta_zone)
            zones_used.append(self._zone_manifest("logo", "logo", logo_box))
            return image, {
                "logo_rendered": logo_rendered,
                "image_rendered": image_rendered,
                "template_rendered": False,
                "image_assessments": image_assessments,
                "render_path": "social_page",
                "layout_variant": "linkedin_insight_panel",
                "zones_used": zones_used,
                "text_blocks_used": [
                    {"role": "brand", "text": brand_name},
                    {"role": "section_label", "text": section_label},
                    {"role": "headline", "text": headline_text},
                    {"role": "supporting_line", "text": supporting_line or body_text},
                    {"role": "proof_points", "text": proof_points},
                    {"role": "stat_highlights", "text": stat_highlights},
                    {"role": "cta", "text": cta_text},
                ],
                "text_fit": text_fit_manifest,
                "asset_boxes": asset_boxes,
                "overlap_checks": overlap_checks,
                "pre_shortening": pre_shortening,
                "content_structure_type": page_text.get("content_structure_type") or "static",
            }

        margin_x = max(48, int(size["width"] * 0.055))
        margin_y = max(34, int(size["height"] * 0.07))
        gap = max(26, int(size["width"] * 0.03))
        text_card = (margin_x, margin_y + 10, int(size["width"] * 0.47), size["height"] - margin_y)
        image_box = (text_card[2] + gap, margin_y + 44, size["width"] - margin_x, size["height"] - margin_y - 24)
        cta_box = (text_card[0] + 40, text_card[3] - 104, text_card[0] + 338, text_card[3] - 28)

        self._draw_panel(draw, (text_card[0], text_card[1], text_card[2], text_card[3]), surface, radius=32, outline=stroke, width=2)
        self._draw_panel(draw, (image_box[0] - 14, image_box[1] - 18, image_box[2] + 14, image_box[3] + 14), accent_soft, radius=36)
        self._draw_section_label(draw, section_label, (text_card[0] + 40, text_card[1] + 26, text_card[0] + 220, text_card[1] + 60), primary, accent)

        headline_box = (text_card[0] + 40, text_card[1] + 78, text_card[2] - 34, text_card[1] + 184)
        support_box = (text_card[0] + 40, text_card[1] + 194, text_card[2] - 34, text_card[1] + 248)
        headline_zone = self._zone_manifest("headline", "headline", headline_box, 3)
        support_zone = self._zone_manifest("supporting_line", "supporting_line", support_box, 2)
        cta_zone = self._zone_manifest("cta", "cta", cta_box, 2)

        headline_fit = self._draw_text_block(
            draw,
            headline_text,
            type("Zone", (), headline_zone)(),
            primary,
            46,
            padding=0,
            style_hint=self._zone_style_hint(payload, "headline"),
        )
        support_fit = self._draw_text_block(
            draw,
            supporting_line or body_text,
            type("Zone", (), support_zone)(),
            secondary_text,
            22,
            padding=0,
            style_hint=self._zone_style_hint(payload, "body"),
        )
        text_fit_manifest.extend([headline_fit, support_fit])
        occupied_text_boxes = _occupied_from_fit([headline_fit, support_fit])
        image_box, image_adjusted = self._adapt_box_away_from_reserved(
            image_box,
            reserved_boxes=occupied_text_boxes,
            canvas_width=size["width"],
            canvas_height=size["height"],
            padding=18,
        )
        image_rendered = False
        if payload.image_assets and page_text.get("show_image"):
            visual_result = self._paste_visual_card(
                image,
                payload.image_assets[0].storage_path,
                image_box,
                radius=32,
                frame_fill=surface,
                padding=0,
            )
            image_rendered = bool(visual_result.get("rendered"))
            if visual_result.get("rendered"):
                image_assessments.append(visual_result)
        if not image_rendered:
            self._draw_editorial_motif(draw, image_box, primary, accent, background)
        asset_boxes.append({"role": "image", "box": image_box, "adjusted": image_adjusted})
        overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(image_box, occupied_text_boxes, padding=18)})
        badges_box = (text_card[0] + 40, text_card[1] + 264, text_card[2] - 40, text_card[1] + 308)
        chip_items = stat_highlights or badges
        self._draw_badges(draw, chip_items, badges_box, primary, accent)
        if cta_text:
            text_fit_manifest.append(self._draw_text_block(
                draw,
                cta_text,
                type("Zone", (), cta_zone)(),
                light_text,
                22,
                background_fill=primary,
                padding=16,
            ))
        logo_rendered = False
        logo_box = None
        logo_box = (image_box[2] - 178, image_box[1] - 30, image_box[2] - 24, image_box[1] + 28)
        logo_box, logo_adjusted = self._adapt_box_away_from_reserved(
            logo_box,
            reserved_boxes=occupied_text_boxes,
            canvas_width=size["width"],
            canvas_height=size["height"],
            padding=20,
        )
        logo_rendered = self._render_logo_box(
            canvas=image,
            draw=draw,
            payload=payload,
            box=logo_box,
            primary=primary,
            accent=accent,
            fill=primary,
        )
        asset_boxes.append({"role": "logo", "box": logo_box, "adjusted": logo_adjusted})
        overlap_checks.append({"role": "logo", "passed": not self._box_overlaps_any(logo_box, occupied_text_boxes, padding=20)})
        zones_used = [
            headline_zone,
            support_zone,
            self._zone_manifest("image", "image", image_box),
            cta_zone,
        ]
        zones_used.insert(0, self._zone_manifest("section_label", "section_label", (text_card[0] + 40, text_card[1] + 26, text_card[0] + 220, text_card[1] + 60), 1))
        zones_used.append(self._zone_manifest("logo", "logo", logo_box))
        if chip_items:
            zones_used.append(self._zone_manifest("badges", "badges", badges_box, 1))
        return image, {
            "logo_rendered": logo_rendered,
            "image_rendered": image_rendered,
            "template_rendered": False,
            "image_assessments": image_assessments,
            "render_path": "social_page",
            "layout_variant": "wide_split_editorial",
            "zones_used": zones_used,
            "text_blocks_used": [
                {"role": "brand", "text": brand_name},
                {"role": "section_label", "text": section_label},
                {"role": "headline", "text": headline_text},
                {"role": "supporting_line", "text": supporting_line or body_text},
                {"role": "cta", "text": cta_text},
                {"role": "badges", "text": chip_items},
            ],
            "text_fit": text_fit_manifest,
            "asset_boxes": asset_boxes,
            "overlap_checks": overlap_checks,
            "pre_shortening": pre_shortening,
            "content_structure_type": page_text.get("content_structure_type") or "static",
        }

    def _render_infographic_page(
        self,
        payload: RendererInput,
        size: dict[str, int],
        page_text: dict,
        preset: str,
        background: tuple[int, int, int],
        primary: tuple[int, int, int],
        accent: tuple[int, int, int],
        secondary_text: tuple[int, int, int],
        light_text: tuple[int, int, int],
    ) -> tuple[Image.Image, dict[str, object]]:
        gradient_spec = self._resolve_gradient_spec(payload, background, primary, accent)
        image = self._build_background_canvas(size["width"], size["height"], background, gradient_spec)
        draw = ImageDraw.Draw(image)
        palette = self._resolve_palette_roles(payload)
        color_tokens = self._reference_palette_tokens(
            payload,
            palette=palette,
            background=background,
            primary=primary,
            accent=accent,
        )
        background = color_tokens["background"]
        primary = color_tokens["primary"]
        accent = color_tokens["accent"]
        surface = color_tokens["surface"]
        soft_primary = color_tokens["primary_soft"]
        stroke = color_tokens["stroke"]
        image_frame_fill = color_tokens["image_frame_fill"]
        brand_name = self._brand_name(payload)
        section_label = self._section_label(payload, preset)
        headline_text = str(page_text.get("headline_display") or "").strip() or self._headline_copy(page_text.get("headline", ""), "linkedin")
        body_text = str(page_text.get("body_display") or "").strip() or self._social_body_copy(page_text.get("body", ""), "instagram")
        supporting_line = str(page_text.get("supporting_line") or "").strip() or self._supporting_line(payload, preset, page_text.get("body", ""))
        proof_points = self._page_list_value(page_text.get("proof_points")) or self._proof_points(payload, "infographic", page_text.get("body", ""))
        stat_highlights = self._page_list_value(page_text.get("stat_highlights")) or self._stat_highlights(payload, "infographic")
        cta_text = str(page_text.get("cta_display") or page_text.get("cta", "") or "").strip()
        image_assessments: list[dict[str, object]] = []
        text_fit_manifest: list[dict[str, object]] = []
        asset_boxes: list[dict[str, object]] = []
        overlap_checks: list[dict[str, object]] = []
        pre_shortening = {
            "headline": {
                "original_length": len(str(page_text.get("headline", "") or "")),
                "fitted_length": len(headline_text),
                "shortened": str(page_text.get("headline", "") or "").strip() != headline_text.strip(),
            },
            "body": {
                "original_length": len(str(page_text.get("body", "") or "")),
                "fitted_length": len(body_text),
                "shortened": str(page_text.get("body", "") or "").strip() != body_text.strip(),
            },
        }

        margin_x = max(56, int(size["width"] * 0.065))
        margin_y = max(68, int(size["height"] * 0.045))
        card_box = (margin_x, margin_y + 36, size["width"] - margin_x, size["height"] - margin_y)
        hero_box = (card_box[0] + 34, card_box[1] + 284, card_box[2] - 34, card_box[1] + 820)
        headline_box = (card_box[0] + 34, card_box[1] + 76, card_box[2] - 34, card_box[1] + 216)
        support_box = (card_box[0] + 34, card_box[1] + 214, card_box[2] - 34, card_box[1] + 280)
        proof_box = (card_box[0] + 34, hero_box[3] + 42, card_box[2] - 34, hero_box[3] + 270)
        stats_box = (card_box[0] + 34, proof_box[3] + 22, card_box[2] - 34, proof_box[3] + 112)
        cta_box = (card_box[0] + 34, stats_box[3] + 24, card_box[0] + 360, stats_box[3] + 100)

        self._draw_panel(draw, (margin_x + 16, margin_y - 4, size["width"] - margin_x - 140, margin_y + 92), self._blend_color(background, accent, 0.12), radius=30)
        self._draw_panel(draw, card_box, surface, radius=42, outline=stroke, width=2)
        self._draw_section_label(draw, section_label, (card_box[0] + 34, card_box[1] + 28, card_box[0] + 220, card_box[1] + 64), primary, accent)

        logo_rendered = False
        logo_box = (card_box[2] - 264, card_box[1] + 24, card_box[2] - 34, card_box[1] + 84)
        logo_rendered = self._render_logo_box(
            canvas=image,
            draw=draw,
            payload=payload,
            box=logo_box,
            primary=primary,
            accent=accent,
            fill=primary,
        )

        headline_zone = self._zone_manifest("headline", "headline", headline_box, 4)
        support_zone = self._zone_manifest("supporting_line", "supporting_line", support_box, 2)
        headline_fit = self._draw_text_block(
            draw,
            headline_text,
            type("Zone", (), headline_zone)(),
            primary,
            58,
            padding=0,
            style_hint=self._zone_style_hint(payload, "headline"),
        )
        support_fit = self._draw_text_block(
            draw,
            supporting_line or body_text,
            type("Zone", (), support_zone)(),
            secondary_text,
            26,
            padding=0,
            style_hint=self._zone_style_hint(payload, "body"),
        )
        text_fit_manifest.extend([headline_fit, support_fit])
        occupied_text_boxes = [
            tuple(item["occupied_box"])
            for item in (headline_fit, support_fit)
            if item.get("occupied_box")
        ]

        image_rendered = False
        hero_box, hero_adjusted = self._adapt_box_away_from_reserved(
            hero_box,
            reserved_boxes=occupied_text_boxes,
            canvas_width=size["width"],
            canvas_height=size["height"],
            padding=18,
        )
        if payload.image_assets and page_text.get("show_image"):
            visual_result = self._paste_visual_card(
                image,
                payload.image_assets[0].storage_path,
                hero_box,
                radius=34,
                frame_fill=image_frame_fill,
                padding=0,
            )
            image_rendered = bool(visual_result.get("rendered"))
            if visual_result.get("rendered"):
                image_assessments.append(visual_result)
        if not image_rendered:
            self._draw_editorial_motif(draw, hero_box, primary, accent, background)
        asset_boxes.append({"role": "image", "box": hero_box, "adjusted": hero_adjusted})
        overlap_checks.append({"role": "image", "passed": not self._box_overlaps_any(hero_box, occupied_text_boxes, padding=18)})

        # Extract badge_style from scene graph if available
        badge_style = None
        if payload.scene_graph and hasattr(payload.scene_graph, 'elements'):
            proof_element = next(
                (el for el in payload.scene_graph.elements if el.role == "proof_points"),
                None
            )
            if proof_element and proof_element.validation_hints:
                badge_style = proof_element.validation_hints.get("badge_style")

        proof_zones = self._draw_bullet_list(draw, proof_points, proof_box, secondary_text, accent, base_size=23, badge_style=badge_style)
        stat_zones = self._draw_stat_cards(draw, stat_highlights, stats_box, primary, accent, background)
        occupied_text_boxes.extend(
            (
                int(zone["x"]),
                int(zone["y"]),
                int(zone["x"]) + int(zone["width"]),
                int(zone["y"]) + int(zone["height"]),
            )
            for zone in [*proof_zones, *stat_zones]
        )
        if cta_text:
            text_fit_manifest.append(self._draw_text_block(
                draw,
                cta_text,
                type("Zone", (), self._zone_manifest("cta", "cta", cta_box, 2))(),
                light_text,
                24,
                background_fill=primary,
                padding=18,
            ))

        zones_used = [
            self._zone_manifest("logo", "logo", logo_box),
            headline_zone,
            support_zone,
            self._zone_manifest("image", "image", hero_box),
            *proof_zones,
            *stat_zones,
        ]
        if cta_text:
            zones_used.append(self._zone_manifest("cta", "cta", cta_box, 2))

        return image, {
            "logo_rendered": logo_rendered,
            "image_rendered": image_rendered,
            "template_rendered": False,
            "image_assessments": image_assessments,
            "render_path": "social_page",
            "layout_variant": "infographic_storyboard",
            "zones_used": zones_used,
            "text_blocks_used": [
                {"role": "brand", "text": brand_name},
                {"role": "section_label", "text": section_label},
                {"role": "headline", "text": headline_text},
                {"role": "supporting_line", "text": supporting_line or body_text},
                {"role": "proof_points", "text": proof_points},
                {"role": "stat_highlights", "text": stat_highlights},
                {"role": "cta", "text": cta_text},
            ],
            "text_fit": text_fit_manifest,
            "asset_boxes": asset_boxes,
            "overlap_checks": overlap_checks,
            "pre_shortening": pre_shortening,
            "content_structure_type": page_text.get("content_structure_type") or "infographic",
        }

    @staticmethod
    def _images_to_pdf_bytes(images: list[Image.Image]) -> bytes:
        if not images:
            return b""
        buffer = BytesIO()
        base = images[0].convert("RGB")
        rest = [page.convert("RGB") for page in images[1:]]
        base.save(buffer, format="PDF", save_all=True, append_images=rest)
        return buffer.getvalue()

    def _assess_render_quality(
        self,
        *,
        payload: RendererInput,
        render_flags: list[dict[str, object]],
        size: dict[str, int],
        decorative_rendered: bool,
    ) -> dict[str, object]:
        page_count = len(render_flags)
        template_rendered = any(bool(flag.get("template_rendered")) for flag in render_flags)
        image_rendered = any(bool(flag.get("image_rendered")) for flag in render_flags)
        logo_rendered = any(bool(flag.get("logo_rendered")) for flag in render_flags)
        issues: list[str] = []
        image_assessments = [
            assessment
            for flag in render_flags
            for assessment in (flag.get("image_assessments") or [])
            if isinstance(assessment, dict)
        ]
        image_fit_score = (
            sum(float(assessment.get("score") or 0.0) for assessment in image_assessments) / len(image_assessments)
            if image_assessments
            else None
        )
        contain_fallback_used = any(str(assessment.get("fit_mode") or "") == "contain" for assessment in image_assessments)
        text_fit_items = [
            item
            for flag in render_flags
            for item in (flag.get("text_fit") or [])
            if isinstance(item, dict)
        ]
        overlap_results = [
            bool(item.get("passed"))
            for flag in render_flags
            for item in (flag.get("overlap_checks") or [])
            if isinstance(item, dict)
        ]
        pre_shortening_items = [
            item
            for flag in render_flags
            for item in [(flag.get("pre_shortening") or {})]
            if isinstance(item, dict)
        ]
        truncated_count = sum(1 for item in text_fit_items if bool(item.get("truncated")))
        min_font_size = min(
            (int(item.get("font_size") or 999) for item in text_fit_items if item.get("font_size") is not None),
            default=None,
        )
        failed_overlap_count = sum(1 for passed in overlap_results if not passed)
        pre_shortened_count = sum(
            1
            for entry in pre_shortening_items
            for item in entry.values()
            if isinstance(item, dict) and bool(item.get("shortened"))
        )
        ellipsis_count = sum(
            1
            for flag in render_flags
            for block in (flag.get("text_blocks_used") or [])
            if isinstance(block, dict) and "..." in str(block.get("text") or "")
        )

        brand_score = 0.58
        if logo_rendered:
            brand_score += 0.14
        elif payload.logo_asset_path:
            brand_score -= 0.08
            issues.append("logo_asset_not_rendered")
        if self._resolve_palette_roles(payload):
            brand_score += 0.14
        if decorative_rendered:
            brand_score += 0.07
        if payload.font_asset_paths:
            brand_score += 0.07

        fidelity_score = 0.66
        style_reference_only = self._template_is_style_reference_only(payload)
        if payload.blueprint.source_mode == "exact_template":
            fidelity_score = 0.9 if template_rendered else 0.28
            if not template_rendered:
                issues.append("template_background_missing")
        elif payload.blueprint.source_mode == "adapted_template":
            if style_reference_only and not template_rendered:
                fidelity_score = 0.78 if image_rendered else 0.6
            else:
                fidelity_score = 0.82 if template_rendered else 0.42
            if not template_rendered and not style_reference_only:
                issues.append("template_adaptation_missing")
        else:
            fidelity_score = 0.8 if render_flags and str(render_flags[0].get("layout_variant", "")).strip() else 0.45

        richness_score = 0.46
        if image_rendered:
            richness_score += 0.24
        if decorative_rendered:
            richness_score += 0.12
        if payload.blueprint.source_mode == "synthesized_layout":
            richness_score += 0.08
        if not image_rendered and not decorative_rendered and payload.blueprint.source_mode == "synthesized_layout":
            issues.append("synthesized_visual_relies_on_motif_only")

        readability_score = 0.72
        if page_count > 1:
            readability_score += 0.04
        if size.get("width", 0) < 900:
            readability_score -= 0.04
        if pre_shortened_count:
            readability_score -= min(pre_shortened_count * 0.03, 0.12)
            issues.append("copy_shortened_before_fit")
        if truncated_count:
            readability_score -= min(truncated_count * 0.08, 0.24)
            issues.append("text_truncation_detected")
        if ellipsis_count:
            readability_score -= min(ellipsis_count * 0.05, 0.15)
            issues.append("ellipsis_visible")
        if failed_overlap_count:
            readability_score -= min(failed_overlap_count * 0.12, 0.3)
            richness_score -= min(failed_overlap_count * 0.06, 0.12)
            issues.append("text_visual_overlap_detected")
        if min_font_size is not None and min_font_size < 14:
            readability_score -= 0.1
            issues.append("font_too_small")
        if image_fit_score is not None and image_fit_score < 0.56:
            readability_score -= 0.06
            richness_score -= 0.05
            issues.append("image_crop_risk")
        elif contain_fallback_used:
            richness_score -= 0.02
            issues.append("image_fit_contained")

        visual_rules = payload.brand_visual_rules if isinstance(payload.brand_visual_rules, dict) else {}
        premium_craft_score = 0.5
        if isinstance(visual_rules.get("visual_craft"), dict) or isinstance(visual_rules.get("composition_logic"), dict):
            premium_craft_score += 0.06
        if image_rendered:
            premium_craft_score += 0.16
        if decorative_rendered:
            premium_craft_score += 0.08
        if template_rendered:
            premium_craft_score += 0.08
        if payload.font_asset_paths:
            premium_craft_score += 0.04
        if payload.blueprint.source_mode == "exact_template":
            premium_craft_score += 0.05
        elif payload.blueprint.source_mode == "adapted_template":
            premium_craft_score += 0.03
        if style_reference_only and not template_rendered:
            premium_craft_score -= 0.08
        if contain_fallback_used:
            premium_craft_score -= 0.04
        if image_fit_score is not None and image_fit_score < 0.56:
            premium_craft_score -= 0.06
        if failed_overlap_count:
            premium_craft_score -= min(failed_overlap_count * 0.04, 0.12)
        if not image_rendered and not decorative_rendered:
            premium_craft_score -= 0.12
        premium_craft_score = max(0.0, premium_craft_score)
        if premium_craft_score < 0.56:
            issues.append("premium_visual_craft_weak")

        overall_score = round(
            min(
                1.0,
                (brand_score * 0.26)
                + (fidelity_score * 0.24)
                + (richness_score * 0.18)
                + (readability_score * 0.14)
                + (premium_craft_score * 0.18),
            ),
            3,
        )
        if overall_score < 0.68:
            issues.append("review_recommended")
        return {
            "overall_score": overall_score,
            "brand_fidelity_score": round(min(brand_score, 1.0), 3),
            "fidelity_score": round(min(fidelity_score, 1.0), 3),
            "visual_richness_score": round(min(richness_score, 1.0), 3),
            "readability_score": round(min(readability_score, 1.0), 3),
            "premium_craft_score": round(min(premium_craft_score, 1.0), 3),
            "image_fit_score": round(image_fit_score, 3) if image_fit_score is not None else None,
            "text_truncation_count": truncated_count,
            "failed_overlap_count": failed_overlap_count,
            "pre_shortened_block_count": pre_shortened_count,
            "min_font_size": min_font_size,
            "issues": issues,
        }

    async def render(self, payload: RendererInput) -> RendererResponse:
        started_at = perf_counter()
        previous_payload = self.payload
        self.payload = payload
        typography = (payload.brand_visual_rules or {}).get("typography", {}) if payload.brand_visual_rules else {}
        self._active_font_candidates = [
            Path(self.storage.absolute_path(path))
            for path in payload.font_asset_paths
            if str(path).strip()
        ]
        self._active_font_bindings = [
            {
                "family_name": str(binding.get("family_name") or ""),
                "storage_path": str(binding.get("storage_path") or ""),
            }
            for binding in (typography.get("uploaded_font_assets", []) or [])
            if isinstance(binding, dict) and str(binding.get("storage_path") or "").strip()
        ]
        self._used_font_paths = set()
        self._used_font_families = set()
        self._requested_font_families = set()
        try:
            template_sizing = (payload.template_metadata or {}).get("sizing_rules") or {}
            size = payload.studio_panel.get("size") or {
                "width": template_sizing.get("width", self.settings.renderer_default_width),
                "height": template_sizing.get("height", self.settings.renderer_default_height),
            }
            if "width" not in size or "height" not in size:
                size = {
                    "width": size.get("width", template_sizing.get("width", self.settings.renderer_default_width)),
                    "height": size.get("height", template_sizing.get("height", self.settings.renderer_default_height)),
            }
            rendered_pages: list[Image.Image] = []
            render_flags: list[dict[str, object]] = []
            page_render_durations_ms: list[float] = []
            if self._should_render_scene_graph_direct(payload):
                page_started_at = perf_counter()
                page_image, flags = self._render_scene_graph(payload, size)
                rendered_pages.append(page_image)
                render_flags.append(flags)
                page_render_durations_ms.append(round((perf_counter() - page_started_at) * 1000, 2))
            else:
                page_payloads = self._build_page_payloads(payload)
                for page_payload in page_payloads:
                    page_started_at = perf_counter()
                    page_image, flags = self._render_page(payload, size, page_payload)
                    reserved_boxes = [
                        tuple(item["occupied_box"])
                        for item in (flags.get("text_fit") or [])
                        if isinstance(item, dict) and item.get("occupied_box")
                    ]
                    decorative_zones = self._apply_decorative_assets(page_image, payload, reserved_boxes=reserved_boxes)
                    if decorative_zones:
                        flags["decorative_rendered"] = True
                        flags["zones_used"] = [*flags.get("zones_used", []), *decorative_zones]
                    rendered_pages.append(page_image)
                    render_flags.append(flags)
                    page_render_durations_ms.append(round((perf_counter() - page_started_at) * 1000, 2))

            preview_asset = self._save_image_asset(
                payload.tenant_id,
                payload.brand_space_id,
                f"preview-{payload.content_version_id}.png",
                rendered_pages[0],
                "render_preview",
            )

            export_assets: list[dict] = []
            target_type = payload.studio_panel.get("file_type", payload.blueprint.export_format or "png")
            if target_type == "jpg":
                for index, page in enumerate(rendered_pages, start=1):
                    export_assets.append(
                        self._save_image_asset(
                            payload.tenant_id,
                            payload.brand_space_id,
                            f"export-{payload.content_version_id}-p{index}.jpg",
                            page,
                            "render_export",
                        )
                    )
            elif target_type == "pdf":
                stored = self.storage.save_bytes(
                    payload.tenant_id,
                    payload.brand_space_id,
                    "generated",
                    f"export-{payload.content_version_id}.pdf",
                    self._images_to_pdf_bytes(rendered_pages),
                )
                export_assets.append(
                    {
                        "asset_id": uuid4(),
                        "mime_type": "application/pdf",
                        "storage_path": stored.storage_path,
                        "width": rendered_pages[0].width,
                        "height": rendered_pages[0].height,
                        "asset_role": "render_export",
                    }
                )
            elif target_type == "doc":
                stored = self._document_export_multi(payload, rendered_pages)
                export_assets.append(
                    {
                        "asset_id": uuid4(),
                        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        "storage_path": stored.storage_path,
                        "width": None,
                        "height": None,
                        "asset_role": "generated_document",
                    }
                )
            else:
                if len(rendered_pages) == 1:
                    export_assets.append(
                        self._save_image_asset(
                            payload.tenant_id,
                            payload.brand_space_id,
                            f"export-{payload.content_version_id}.png",
                            rendered_pages[0],
                            "render_export",
                        )
                    )
                else:
                    for index, page in enumerate(rendered_pages, start=1):
                        export_assets.append(
                            self._save_image_asset(
                                payload.tenant_id,
                                payload.brand_space_id,
                                f"export-{payload.content_version_id}-p{index}.png",
                                page,
                                "render_export",
                            )
                        )

            image_rendered = any(bool(flag["image_rendered"]) for flag in render_flags)
            logo_rendered = any(bool(flag["logo_rendered"]) for flag in render_flags)
            template_rendered = any(bool(flag["template_rendered"]) for flag in render_flags)
            decorative_rendered = any(bool(flag.get("decorative_rendered")) for flag in render_flags)
            primary_manifest = render_flags[0] if render_flags else {}
            quality_assessment = self._assess_render_quality(
                payload=payload,
                render_flags=render_flags,
                size=size,
                decorative_rendered=decorative_rendered,
            )
            if payload.blueprint.source_mode in {"exact_template", "adapted_template"} and not template_rendered and not self._template_is_style_reference_only(payload):
                logger.warning(
                    "renderer.render.template_not_applied brand_space_id=%s content_version_id=%s source_mode=%s source_template_id=%s layout_variant=%s",
                    payload.brand_space_id,
                    payload.content_version_id,
                    payload.blueprint.source_mode,
                    payload.blueprint.source_template_id,
                    primary_manifest.get("layout_variant", "generic_zones"),
                )
            return RendererResponse(
                preview_asset=preview_asset,
                export_assets=export_assets,
                renderer_metadata={
                    "deterministic": True,
                    "overflow_strategy": payload.blueprint.overflow_strategy,
                    "canvas_size": size,
                    "logo_rendered": logo_rendered,
                    "image_rendered": image_rendered,
                    "template_rendered": template_rendered,
                    "decorative_rendered": decorative_rendered,
                    "layout_type": payload.blueprint.layout_type,
                    "source_mode": (payload.scene_graph.layout_mode if payload.scene_graph else payload.blueprint.source_mode),
                    "source_template_id": (
                        str((payload.creative_decision or {}).get("selected_template_id") or "")
                        or (payload.blueprint.source_template_id if payload.blueprint else None)
                    ),
                    "adaptation_plan": (
                        payload.scene_graph.template_adaptation
                        if payload.scene_graph
                        else payload.blueprint.adaptation_plan
                    ),
                    "brand_rules_applied": payload.blueprint.brand_rules_applied if payload.blueprint else {},
                    "layout_decision": payload.layout_decision,
                    "creative_decision": payload.creative_decision,
                    "validation_report": payload.validation_report,
                    "page_count": len(rendered_pages),
                    "layout_variant": primary_manifest.get("layout_variant", "generic_zones"),
                    "content_structure_type": primary_manifest.get("content_structure_type") or payload.blueprint.layout_type,
                    "latency_ms": {
                        "render_total_ms": round((perf_counter() - started_at) * 1000, 2),
                        "page_render_ms": page_render_durations_ms,
                    },
                    "quality_assessment": quality_assessment,
                    "render_path": primary_manifest.get("render_path", "page"),
                    "image_fit": {
                        "assessments": [
                            assessment
                            for flag in render_flags
                            for assessment in (flag.get("image_assessments") or [])
                            if isinstance(assessment, dict)
                        ],
                    },
                    "render_manifest": {
                        "zones_used": primary_manifest.get("zones_used", []),
                        "text_blocks_used": primary_manifest.get("text_blocks_used", []),
                        "text_fit": primary_manifest.get("text_fit", []),
                        "asset_boxes": primary_manifest.get("asset_boxes", []),
                        "overlap_checks": primary_manifest.get("overlap_checks", []),
                        "pre_shortening": primary_manifest.get("pre_shortening", {}),
                        "content_structure_type": primary_manifest.get("content_structure_type") or payload.blueprint.layout_type,
                    },
                    "font_resolution": {
                        "available_uploaded_fonts": [
                            str(binding.get("family_name") or "")
                            for binding in self._active_font_bindings
                            if str(binding.get("family_name") or "").strip()
                        ],
                        "requested_font_families": sorted(self._requested_font_families),
                        "used_font_families": sorted(self._used_font_families),
                        "used_font_paths": sorted(self._used_font_paths),
                    },
                    "asset_resolution": {
                        "image_asset_paths": [asset.storage_path for asset in payload.image_assets],
                        "decorative_asset_paths": [asset.storage_path for asset in payload.decorative_assets],
                        "logo_asset_path": payload.logo_asset_path,
                        "template_asset_path": payload.template_asset_path,
                    },
                },
            )
        finally:
            self.payload = previous_payload
            self._active_font_candidates = []
            self._active_font_bindings = []
            self._used_font_paths = set()
            self._used_font_families = set()
            self._requested_font_families = set()
